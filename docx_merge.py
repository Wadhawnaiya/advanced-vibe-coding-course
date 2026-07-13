"""Shared toolkit for assembling the branded, consolidated Vibe Coding documents."""
import json
import os
import re

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
BRAND_PATH = os.path.join(BASE_DIR, "brand.json")

PART_LETTER_RE = re.compile(r"^Part [A-Z] — (.+)$")

# Matches a source document's own closing/sign-off marker, e.g.
# "— End of Master Teaching Guide (Edition 2.0) —", "End of Course Guide — ..."
# (the leading em dash is not always present), or "Use with: <old filenames>".
_CLOSING_TRAILER_RE = re.compile(r"^(—\s*)?End of\b|^Use with:", re.IGNORECASE)


def load_brand():
    with open(BRAND_PATH, encoding="utf-8") as f:
        return json.load(f)


def _color(hexstr):
    return RGBColor(int(hexstr[0:2], 16), int(hexstr[2:4], 16), int(hexstr[4:6], 16))


def _set_run(run, size=11, bold=False, italic=False, color=None, font="Calibri"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = _color(color)


def load_excerpt(path, start_heading, end_heading=None):
    """Open `path` fresh and trim its body to the paragraph range
    [start_heading, end_heading). Matching is exact on stripped paragraph
    text. Raises AssertionError if start_heading is not found."""
    doc = Document(path)
    body = doc.element.body
    start_el = None
    end_el = None
    for para in doc.paragraphs:
        text = para.text.strip()
        if start_el is None and text == start_heading:
            start_el = para._p
            continue
        if start_el is not None and end_heading and text == end_heading:
            end_el = para._p
            break
    assert start_el is not None, f"heading not found in {path}: {start_heading!r}"
    children = list(body)
    start_idx = children.index(start_el)
    end_idx = children.index(end_el) if end_el is not None else len(children) - 1
    keep_ids = {id(c) for c in children[start_idx:end_idx]}
    for child in children:
        if id(child) not in keep_ids:
            body.remove(child)
    return doc


def strip_closing_trailer(doc, tail_window=15):
    """Remove a source document's own closing/sign-off paragraphs (e.g.
    '— End of X —', 'End of X — ...', 'Use with: <old filenames>', and any
    stray sign-off line that immediately follows one) from the tail of
    `doc`'s body. Used on excerpts loaded with `end_heading=None`, which run
    to the true end of their source document and therefore also pull in
    that source's own trailer.

    A naive "scan backward from the end, stop at the first non-matching
    paragraph" approach is not reliable here: some sources append a short
    sign-off *after* their closing marker (e.g. a "Primary references: ..."
    or tagline paragraph) that doesn't itself match the trailer pattern, so
    stopping at the first non-match would miss the marker entirely. Instead
    this scans the last `tail_window` paragraphs for the *earliest*
    paragraph matching the trailer pattern and removes it and everything
    after it. If no match is found in that window, nothing is removed.
    """
    body = doc.element.body
    children = list(body)
    sectpr = children[-1] if children and children[-1].tag.endswith("}sectPr") else None
    content = children[:-1] if sectpr is not None else children

    window_start = max(0, len(content) - tail_window)
    cut_at = None
    for i in range(window_start, len(content)):
        text = "".join(content[i].itertext()).strip()
        if text and _CLOSING_TRAILER_RE.match(text):
            cut_at = i
            break
    if cut_at is None:
        return doc
    for el in content[cut_at:]:
        body.remove(el)
    return doc


def _style_by_name(doc, name):
    """Look up a style by its resolved UI name.

    Some source documents (built by a non-python-docx docx generator) store the
    style's raw `w:name` value already in UI form (e.g. literally "Heading 2")
    rather than Word's conventional internal lowercase form ("heading 2"). In
    that case `doc.styles[name]` raises KeyError because python-docx's BabelFish
    translation looks for the lowercase form and doesn't find it, even though
    the style is clearly present. Iterating and comparing the resolved `.name`
    (the same accessor used to *read* style names elsewhere in this module)
    works for both conventions.
    """
    for style in doc.styles:
        if style.name == name:
            return style
    raise KeyError(f"no style with name {name!r}")


def demote_part_headings(doc):
    """Turn source 'Part X — Title' H1s into H2s with the 'Part X — ' prefix
    stripped, so they nest under one umbrella H1 in the merged document."""
    for para in doc.paragraphs:
        if para.style is None or para.style.name != "Heading 1":
            continue
        m = PART_LETTER_RE.match(para.text.strip())
        if not m:
            continue
        new_text = m.group(1)
        for i, run in enumerate(para.runs):
            run.text = new_text if i == 0 else ""
        para.style = _style_by_name(doc, "Heading 2")
    return doc


def recolor_headings(doc, brand):
    """Force every Heading 1/2/3 paragraph's run colors to the brand palette,
    overriding whatever color the source document originally used.

    Excerpts merged in via docxcompose preserve each source's original
    run-level formatting verbatim, including any heading colors that predate
    the brand's unified green palette (e.g. the Full Course Guide's blue
    1B4F72/2874A6). This is a single whole-document pass over the final
    composed `doc`, run once at the end of assembly, so it recolors every
    heading regardless of which source excerpt it came from."""
    palette = brand["palette"]
    level_colors = {
        "Heading 1": palette["primary"],
        "Heading 2": palette["accent"],
        "Heading 3": palette["dark"],
    }
    for para in doc.paragraphs:
        if para.style is None:
            continue
        color_hex = level_colors.get(para.style.name)
        if color_hex is None:
            continue
        for run in para.runs:
            run.font.color.rgb = _color(color_hex)
    return doc


def new_base_document():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    return doc


def render_cover(doc, brand, light=False):
    palette = brand["palette"]
    for _ in range(3):
        doc.add_paragraph()
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run(t.add_run(brand["course"]["title"].upper()), size=26, bold=True, color=palette["primary"])
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run(t.add_run(brand["course"]["subtitle"]), size=13, italic=True, color=palette["muted"])
    doc.add_paragraph()
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run(t.add_run(brand["name"]), size=18, bold=True, color=palette["dark"])
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run(t.add_run(brand["credentials_line"]), size=11, color=palette["muted"])
    doc.add_paragraph()
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run(t.add_run(brand["tagline"]), size=11, italic=True, color=palette["accent"])
    doc.add_paragraph()
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run(t.add_run(brand["orgs"]["provider_line"]), size=10, color=palette["muted"])
    if not light:
        t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact = brand["contact"]
        _set_run(
            t.add_run(f"{contact['phone']}  ·  {contact['email']}  ·  {contact['linkedin']}  ·  {contact['website']}"),
            size=9, color=palette["muted"],
        )
    doc.add_paragraph()
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run(t.add_run(f"Version {brand['course']['version']} · {brand['course']['date']}"), size=9, color=palette["muted"])
    doc.add_page_break()
    return doc


def render_about_page(doc, brand, light=False):
    palette = brand["palette"]
    h = doc.add_heading("About the Instructor" if not light else "Course Provider", level=1)
    for run in h.runs:
        run.font.color.rgb = _color(palette["primary"])
    p1 = doc.add_paragraph()
    _set_run(p1.add_run(brand["name"]), size=14, bold=True, color=palette["dark"])
    p2 = doc.add_paragraph()
    _set_run(p2.add_run(brand["credentials_line"]), size=11, color=palette["muted"])
    doc.add_paragraph()
    body = doc.add_paragraph()
    _set_run(body.add_run(brand["bio"] if not light else brand["tagline"]), size=11)
    if not light:
        h2 = doc.add_heading("Roles", level=2)
        for run in h2.runs:
            run.font.color.rgb = _color(palette["accent"])
        for role in brand["roles"]:
            b = doc.add_paragraph(style="List Bullet")
            _set_run(b.add_run(role), size=11)
        h3 = doc.add_heading("Credentials", level=2)
        for run in h3.runs:
            run.font.color.rgb = _color(palette["accent"])
        for cred in brand["credentials_detail"]:
            b = doc.add_paragraph(style="List Bullet")
            _set_run(b.add_run(cred), size=11)
        h4 = doc.add_heading("Contact", level=2)
        for run in h4.runs:
            run.font.color.rgb = _color(palette["accent"])
        contact = brand["contact"]
        for label, value in [
            ("Phone", contact["phone"]),
            ("Email", contact["email"]),
            ("LinkedIn", contact["linkedin"]),
            ("Website", contact["website"]),
            ("Location", contact["location"]),
        ]:
            b = doc.add_paragraph()
            _set_run(b.add_run(f"{label}: "), size=11, bold=True)
            _set_run(b.add_run(value), size=11)
    doc.add_page_break()
    return doc


def apply_header_footer(doc, brand, subtitle):
    palette = brand["palette"]
    section = doc.sections[0]
    header_p = section.header.paragraphs[0]
    header_p.text = ""
    _set_run(header_p.add_run(f"{brand['course']['title']} — {subtitle}"), size=9, color=palette["muted"])
    footer_p = section.footer.paragraphs[0]
    footer_p.text = ""
    _set_run(footer_p.add_run(f"{brand['name']} · {brand['orgs']['primary']}"), size=9, color=palette["muted"])
    return doc
