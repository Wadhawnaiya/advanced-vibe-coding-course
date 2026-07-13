#!/usr/bin/env python3
"""
Vibe Coding with Claude Code — MASTER TEACHING GUIDE
Deep, explanation-first curriculum for non-technical learners.
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

OUT = "/home/shailesh/Downloads/Training/vibe-coding/Vibe_Coding_Master_Teaching_Guide_v2.docx"

# Colors
PRIMARY = RGBColor(0x0B, 0x3D, 0x2E)
ACCENT = RGBColor(0x1A, 0x7A, 0x5C)
DARK = RGBColor(0x1C, 0x28, 0x33)
MUTED = RGBColor(0x5D, 0x6D, 0x7E)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

def set_run(run, size=11, bold=False, italic=False, color=DARK, font="Calibri"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color

def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = PRIMARY if level == 1 else ACCENT if level == 2 else DARK
        run.font.name = "Calibri"
    return h

def p(doc, text, size=11, bold=False, italic=False, space_after=8, color=DARK):
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(space_after)
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.line_spacing = 1.15
    run = para.add_run(text)
    set_run(run, size=size, bold=bold, italic=italic, color=color)
    return para

def multi(doc, parts, space_after=8):
    """parts: list of (text, bold, italic, color?)"""
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(space_after)
    para.paragraph_format.line_spacing = 1.15
    for part in parts:
        text = part[0]
        bold = part[1] if len(part) > 1 else False
        italic = part[2] if len(part) > 2 else False
        color = part[3] if len(part) > 3 else DARK
        size = part[4] if len(part) > 4 else 11
        run = para.add_run(text)
        set_run(run, size=size, bold=bold, italic=italic, color=color)
    return para

def bullets(doc, items, size=11):
    for item in items:
        para = doc.add_paragraph(style="List Bullet")
        para.paragraph_format.space_after = Pt(4)
        para.clear()
        run = para.add_run(item)
        set_run(run, size=size)

def numbered(doc, items, size=11):
    for item in items:
        para = doc.add_paragraph(style="List Number")
        para.paragraph_format.space_after = Pt(4)
        para.clear()
        run = para.add_run(item)
        set_run(run, size=size)

def callout(doc, title, body_lines, fill="E8F6F1"):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.rows[0].cells[0]
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill}"/>')
    cell._tePr = cell._tc.get_or_add_tcPr()
    cell._tc.get_or_add_tcPr().append(shading)
    cell.text = ""
    p1 = cell.paragraphs[0]
    r1 = p1.add_run(title)
    set_run(r1, size=11, bold=True, color=PRIMARY)
    for line in body_lines:
        pp = cell.add_paragraph()
        rr = pp.add_run(line)
        set_run(rr, size=10, color=DARK)
    doc.add_paragraph()

def code_block(doc, lines):
    if isinstance(lines, str):
        lines = lines.split("\n")
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1E2A32"/>')
    cell._tc.get_or_add_tcPr().append(shading)
    cell.text = ""
    first = True
    for line in lines:
        if first:
            para = cell.paragraphs[0]
            first = False
        else:
            para = cell.add_paragraph()
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.space_before = Pt(0)
        run = para.add_run(line if line else " ")
        set_run(run, size=9, color=RGBColor(0xE8, 0xF0, 0xF2), font="Consolas")
    doc.add_paragraph()

def table_data(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        set_run(run, size=10, bold=True, color=WHITE)
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="0B3D2E"/>')
        cell._tc.get_or_add_tcPr().append(shading)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            set_run(run, size=9)
            if ri % 2 == 0:
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F4F6F7"/>')
                cell._tc.get_or_add_tcPr().append(shading)
    doc.add_paragraph()

def page_break(doc):
    doc.add_page_break()

def setup_styles(doc):
    styles = doc.styles
    # Normal
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = DARK

def build():
    doc = Document()
    setup_styles(doc)
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    # ========== COVER ==========
    for _ in range(4):
        doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("MASTER TEACHING GUIDE")
    set_run(r, size=14, bold=True, color=ACCENT)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Vibe Coding with Claude Code")
    set_run(r, size=28, bold=True, color=PRIMARY)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("From Absolute Zero to Shipping Full-Stack Applications")
    set_run(r, size=14, color=DARK)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Deep explanations · Real features · Real commands · Real fundamentals")
    set_run(r, size=11, italic=True, color=MUTED)

    doc.add_paragraph()
    callout(doc, "What makes this guide different", [
        "This is not a slogan-only “non-tech philosophy” booklet.",
        "Every major tool, command, language concept, and workflow is explained in plain language,",
        "with why it exists, how it works, when you use it, and what can go wrong.",
        "Written so a complete beginner can understand — and so an instructor can teach with confidence.",
    ])

    p(doc, "Primary tool: Claude Code CLI (Anthropic)  ·  Companion skills: Terminal, Git, GitHub, HTML, CSS, JavaScript, Python, Full-stack apps", size=10, color=MUTED)
    p(doc, "Edition 2.0 — Explanation-first rewrite  ·  Grounded in Claude Code official documentation (2026)", size=10, color=MUTED)

    page_break(doc)

    # ========== TOC ==========
    add_heading_styled(doc, "How This Book Is Organized", 1)
    p(doc, "Read in order if you are a complete beginner. Jump by part if you already know one area.")
    toc = [
        ("Part 1", "Mindset & What You Are Actually Learning"),
        ("Part 2", "Computer Foundations (files, folders, internet, accounts)"),
        ("Part 3", "The Terminal — Your Text Remote Control (full command explanations)"),
        ("Part 4", "What Claude Code Is and How It Really Works"),
        ("Part 5", "Claude Code Features Encyclopedia"),
        ("Part 6", "Slash Commands Encyclopedia (what / why / when / example)"),
        ("Part 7", "CLAUDE.md, Skills, Permissions, Plan Mode, MCP"),
        ("Part 8", "How to Talk to Claude Code (prompt craft that actually works)"),
        ("Part 9", "How the Web Works + Full-Stack Mental Model"),
        ("Part 10", "HTML — Structure of Every Webpage (proper beginner course)"),
        ("Part 11", "CSS — Making Pages Look Right"),
        ("Part 12", "JavaScript — Making Pages Do Things"),
        ("Part 13", "Python — Backend Language Foundations"),
        ("Part 14", "Git — Version Control Explained Properly"),
        ("Part 15", "GitHub — Cloud Collaboration & Portfolio Guide"),
        ("Part 16", "Build Path: Projects from Hello World to Full-Stack"),
        ("Part 17", "Deploying So the World Can Use Your App"),
        ("Part 18", "Capstone, Assessment, and Teaching Notes"),
        ("Appendices", "Glossaries, cheatsheets, troubleshooting, install references"),
    ]
    table_data(doc, ["Part", "What you will master"], toc)
    page_break(doc)

    # ========== PART 1 ==========
    add_heading_styled(doc, "Part 1 — Mindset & What You Are Actually Learning", 1)

    add_heading_styled(doc, "1.1 What “vibe coding” means (precisely)", 2)
    p(doc, "Vibe coding is a way of building software where you describe the outcome you want in natural language, and an AI coding agent (here: Claude Code) explores your project, writes and edits files, runs commands, and iterates — while you stay in control as the decision-maker.")
    p(doc, "It is not magic. Three things are always true:")
    numbered(doc, [
        "The AI can be wrong, incomplete, outdated, or insecure. You must verify.",
        "Clear instructions produce better software than vague wishes.",
        "Understanding the parts of an application (even without memorizing syntax) makes you far more effective at directing the AI.",
    ])
    p(doc, "This course teaches both: how to direct Claude Code professionally, and enough real computer science literacy (HTML, CSS, JavaScript, Python, Git, GitHub, full-stack structure) that you are not helpless when the AI gets stuck.")

    add_heading_styled(doc, "1.2 Your role vs Claude Code’s role", 2)
    table_data(doc, ["You own", "Claude Code owns", "Together you produce"], [
        ["The problem and the user", "Writing most of the code", "A working application"],
        ["What “done” means (acceptance checks)", "Editing multiple files", "Something you can demo"],
        ["Approving dangerous actions", "Running install/test commands", "A Git history of progress"],
        ["Testing like a real user", "Explaining what it changed", "A deployed public link (later)"],
        ["Scope decisions (MVP vs later)", "Suggesting architecture options", "Learning you can reuse forever"],
    ])

    add_heading_styled(doc, "1.3 What “full-stack application” means", 2)
    p(doc, "A full-stack application is software that includes all major layers end-users depend on:")
    bullets(doc, [
        "Frontend — the interface people see in a browser (pages, buttons, forms).",
        "Backend — server-side logic that processes requests, enforces rules, talks to data.",
        "Database (or other storage) — lasting memory for users, posts, orders, etc.",
        "Often authentication — login so people have private data.",
        "Deployment — hosting so the app is available on the internet, not only on your laptop.",
    ])
    p(doc, "By the end of this course you will direct Claude Code to build a small but real full-stack app, put it on GitHub, and deploy it.")

    add_heading_styled(doc, "1.4 Learning contract (read this to students)", 2)
    callout(doc, "We do not pretend code does not exist", [
        "We will not force you to write every line from memory.",
        "We will force you to understand what a line is for, what a folder is for, and what a command does.",
        "If you only click “accept” forever without understanding, you will freeze the first time Claude is wrong.",
        "If you learn concepts + verification + Claude Code craft, you can build for years.",
    ], fill="FEF9E7")
    page_break(doc)

    # ========== PART 2 ==========
    add_heading_styled(doc, "Part 2 — Computer Foundations (Explained Properly)", 1)

    add_heading_styled(doc, "2.1 Files and folders — the real model", 2)
    p(doc, "Everything on a computer is stored as files inside folders (also called directories). A file is a named package of data. A folder is a container that holds files and other folders.")
    p(doc, "Examples of file types you will meet constantly:")
    table_data(doc, ["Extension", "What it usually is", "Why you care"], [
        [".html", "Web page structure", "Browser can open it directly"],
        [".css", "Visual styling rules", "Colors, layout, fonts"],
        [".js", "JavaScript behavior", "Clicks, logic in the browser"],
        [".py", "Python program", "Common backend language"],
        [".json", "Structured data text", "Config and API data"],
        [".md", "Markdown notes/docs", "README, CLAUDE.md"],
        [".env", "Secrets / config", "NEVER publish publicly"],
        [".gitignore", "Git ignore rules", "Keeps secrets out of Git"],
    ])
    p(doc, "A path is the address of a file. Example on Mac/Linux: /Users/alex/Documents/my-app/index.html. Example on Windows: C:\\Users\\alex\\Documents\\my-app\\index.html. You must always know “which folder am I in?” before running tools like Claude Code — because Claude works inside the folder you start it from.")

    add_heading_styled(doc, "2.2 The browser vs the terminal vs the editor", 2)
    bullets(doc, [
        "Browser (Chrome, Edge, Safari): displays websites and web apps. This is where you test the user experience.",
        "Terminal (command line): a text interface to give the computer precise instructions (install tools, run servers, start Claude Code).",
        "Code editor / IDE (VS Code, etc.): a program for viewing and editing many code files with helpful features. Optional early; powerful later.",
        "Claude Code can work mainly from the terminal without you manually editing every file — but you should still open files sometimes to understand them.",
    ])

    add_heading_styled(doc, "2.3 Accounts, passwords, and keys (security literacy)", 2)
    p(doc, "Building apps means creating accounts (Claude, GitHub, hosting). Rules that protect you:")
    numbered(doc, [
        "Use a password manager or unique strong passwords. Never reuse the same password everywhere.",
        "Turn on two-factor authentication (2FA) for GitHub and email when possible.",
        "API keys and tokens are secret passwords for programs. They go in .env files, not in public code, not in screenshots, not in Discord.",
        "If a key leaks, revoke it immediately in the provider’s dashboard and create a new one.",
    ])
    page_break(doc)

    # ========== PART 3 ==========
    add_heading_styled(doc, "Part 3 — The Terminal: Your Text Remote Control", 1)

    add_heading_styled(doc, "3.1 What the terminal actually is", 2)
    p(doc, "The terminal is a program that accepts text commands and runs them on your computer. Before graphical buttons existed, this was how all computers were controlled. Today it is still the fastest, most precise way to:")
    bullets(doc, [
        "Install developer tools",
        "Start web servers",
        "Run tests",
        "Use Git",
        "Launch Claude Code",
    ])
    p(doc, "How to open it:")
    bullets(doc, [
        "macOS: Spotlight (Cmd+Space) → type Terminal. Or open a folder in Finder → right-click → Services / New Terminal at Folder.",
        "Windows: Win key → type PowerShell (preferred for Claude install). Note: a prompt starting with PS means PowerShell. CMD is different.",
        "Linux: Ctrl+Alt+T, or search “Terminal”.",
    ])

    add_heading_styled(doc, "3.2 Anatomy of a command", 2)
    p(doc, "Most commands look like this:")
    code_block(doc, "command  argument1  argument2  --flag  value")
    p(doc, "Example:")
    code_block(doc, "mkdir my-project")
    p(doc, "mkdir is the command (make directory). my-project is the argument (the name of the folder). Press Enter to run. If something fails, the terminal prints an error message — that text is valuable. Copy it. Do not panic.")

    add_heading_styled(doc, "3.3 Essential commands explained one by one", 2)
    p(doc, "Memorize these. You will use them for years.")

    add_heading_styled(doc, "pwd — Print Working Directory", 3)
    p(doc, "What it does: Shows the full path of the folder you are currently inside.")
    p(doc, "Why it matters: Claude Code and many tools act on the current folder. If you are in the wrong place, you create files in the wrong place.")
    code_block(doc, "pwd")

    add_heading_styled(doc, "ls (Mac/Linux) or dir (Windows CMD) — List", 3)
    p(doc, "What it does: Lists files and folders in the current directory.")
    p(doc, "Useful variants:")
    code_block(doc, [
        "ls          # basic list",
        "ls -la      # detailed list including hidden files (Mac/Linux)",
        "dir         # Windows CMD list",
        "Get-ChildItem   # PowerShell list (or alias: ls)",
    ])

    add_heading_styled(doc, "cd — Change Directory", 3)
    p(doc, "What it does: Moves you into another folder.")
    code_block(doc, [
        "cd Documents",
        "cd vibe-coding-course",
        "cd ..          # go up one folder",
        "cd ~           # go to home folder (Mac/Linux)",
        "cd /           # go to system root (careful)",
    ])
    p(doc, "If a folder name has spaces, quote it: cd \"My Projects\"")

    add_heading_styled(doc, "mkdir — Make Directory", 3)
    p(doc, "Creates a new folder.")
    code_block(doc, "mkdir week1-hello")

    add_heading_styled(doc, "touch (Mac/Linux) — Create empty file", 3)
    code_block(doc, "touch notes.txt")
    p(doc, "On Windows PowerShell you can use: New-Item notes.txt")

    add_heading_styled(doc, "clear or cls — Clear the screen", 3)
    p(doc, "Does not delete files. Only clears visual clutter in the terminal window.")
    code_block(doc, [
        "clear    # Mac/Linux",
        "cls      # Windows",
    ])

    add_heading_styled(doc, "cat / type — Show file contents", 3)
    code_block(doc, [
        "cat README.md      # Mac/Linux",
        "type README.md     # Windows CMD",
        "Get-Content README.md  # PowerShell",
    ])

    add_heading_styled(doc, "Commands you must NOT run casually", 3)
    callout(doc, "Danger zone", [
        "rm -rf something  — force-delete folders forever (Mac/Linux). Easy to destroy a project or worse.",
        "del /s or format  — destructive Windows operations.",
        "Never paste random commands from strangers into your terminal.",
        "If Claude suggests a destructive command, read it twice. Prefer safer alternatives.",
    ], fill="FDEDEC")

    add_heading_styled(doc, "3.4 PATH — why “command not found” happens", 2)
    p(doc, "When you type a command, the computer looks for a matching program in a list of folders called PATH. If the program is installed but not on PATH, you get command not found. After installing Claude Code, if claude is not found, you usually need to add ~/.local/bin to PATH (the installer often prints the exact fix). Then open a new terminal window.")
    page_break(doc)

    # ========== PART 4 ==========
    add_heading_styled(doc, "Part 4 — What Claude Code Is and How It Really Works", 1)

    add_heading_styled(doc, "4.1 Definition", 2)
    p(doc, "Claude Code is Anthropic’s agentic coding environment. Unlike a normal chatbot that only replies with text, Claude Code can:")
    bullets(doc, [
        "Read files across your project",
        "Edit and create files",
        "Run shell commands (install packages, start servers, run tests)",
        "Search the codebase",
        "Use extra tools via MCP (Model Context Protocol) when configured",
        "Follow project instructions from CLAUDE.md and skills",
        "Work in a loop: try → observe result → fix → continue",
    ])
    p(doc, "Surfaces (ways to use the same engine): Terminal CLI (this course’s focus), VS Code/Cursor extension, Desktop app, Web (claude.ai/code), JetBrains plugin. Settings like CLAUDE.md and MCP can apply across surfaces.")

    add_heading_styled(doc, "4.2 Account requirements (important)", 2)
    p(doc, "Claude Code requires a paid Claude plan (Pro, Max, Team, or Enterprise) or Anthropic Console / certain cloud provider setups. The free Claude.ai plan alone does not include Claude Code access. Budget for this before teaching a class.")

    add_heading_styled(doc, "4.3 Installation (official methods)", 2)
    p(doc, "System requirements (summary): macOS 13+, Windows 10 1809+, or supported Linux; about 4GB+ RAM; internet required.")
    p(doc, "Recommended native install:")
    p(doc, "macOS / Linux / WSL:", bold=True)
    code_block(doc, "curl -fsSL https://claude.ai/install.sh | bash")
    p(doc, "Windows PowerShell:", bold=True)
    code_block(doc, "irm https://claude.ai/install.ps1 | iex")
    p(doc, "Windows CMD:", bold=True)
    code_block(doc, "curl -fsSL https://claude.ai/install.cmd -o install.cmd && install.cmd && del install.cmd")
    p(doc, "Also available: Homebrew (brew install --cask claude-code), WinGet (winget install Anthropic.ClaudeCode). On Windows, Git for Windows is recommended so Claude can use Bash tooling.")
    p(doc, "Verify:")
    code_block(doc, [
        "claude --version",
        "claude doctor",
    ])
    p(doc, "Start in a project folder:")
    code_block(doc, [
        "cd path/to/your-project",
        "claude",
    ])
    p(doc, "First run opens a browser login. After that, credentials are stored.")

    add_heading_styled(doc, "4.4 The agent loop (how it thinks and acts)", 2)
    p(doc, "When you give Claude Code a task, roughly this happens:")
    numbered(doc, [
        "It reads your message and relevant project context (including CLAUDE.md).",
        "It decides which tools to use: read file, edit file, run command, search, etc.",
        "It may ask permission depending on your permission mode.",
        "It observes the tool results (file contents, command output, errors).",
        "It continues until it believes the task is done — or until you interrupt/redirect.",
    ])
    p(doc, "Critical insight from Anthropic’s best practices: Claude stops when work “looks done.” If you do not give a verification method (run the app, run tests, open the page), mistakes wait for you to notice. Always define what “done” means.")

    add_heading_styled(doc, "4.5 Context window — the hidden resource", 2)
    p(doc, "The context window is the amount of conversation + file content + command output Claude can hold at once. Every message, every file read, and every long log fills it. When it gets full, quality drops: Claude “forgets” earlier instructions or makes more mistakes.")
    p(doc, "Practical habits:")
    bullets(doc, [
        "One major task per session when possible",
        "Use /clear when switching to a completely different task",
        "Use /compact to summarize and free space while staying on the same task",
        "Keep CLAUDE.md short so it does not waste context on fluff",
        "Prefer @filename to point Claude at exact files instead of pasting huge dumps",
    ])
    page_break(doc)

    # ========== PART 5 ==========
    add_heading_styled(doc, "Part 5 — Claude Code Features Encyclopedia", 1)
    p(doc, "This part explains the product features in teaching language. Slash command details are expanded in Part 6.")

    add_heading_styled(doc, "5.1 Interactive session (REPL)", 2)
    p(doc, "When you run claude in a folder, you enter an interactive session. You type English (or a slash command). Claude responds, may use tools, and waits for the next instruction. This is the main classroom mode.")

    add_heading_styled(doc, "5.2 File tools (read / edit / write / search)", 2)
    p(doc, "Claude can open files, modify them, create new ones, and search across the project. You should treat every edit as a proposal: review diffs when possible, especially early in learning. Later you may use acceptEdits mode for speed, but beginners should stay in Manual mode longer.")

    add_heading_styled(doc, "5.3 Bash / shell tool", 2)
    p(doc, "Claude can run terminal commands: npm install, python app.py, git status, pytest, etc. This is why Claude Code feels more powerful than “chat that pastes code.” It can actually execute. Always read commands before approving until you trust the pattern.")

    add_heading_styled(doc, "5.4 Permission modes", 2)
    p(doc, "Permission modes control how often Claude pauses to ask approval. Switch with Shift+Tab in the CLI (cycles Manual → acceptEdits → plan, and more if enabled).")
    table_data(doc, ["Mode", "What it allows without asking (simplified)", "Best for students"], [
        ["Manual (default)", "Reads freely; asks before many edits/commands", "Beginners, sensitive work — START HERE"],
        ["acceptEdits", "Auto-approves many file edits + common file ops", "When reviewing via git diff after"],
        ["plan", "Research + plan; does not edit source", "Before big features — USE OFTEN"],
        ["auto", "Fewer interruptions with safety checks", "Advanced, long tasks"],
        ["bypassPermissions", "Skips checks", "Only isolated VMs/containers — NOT class laptops"],
    ])
    p(doc, "Teaching rule: Plan mode for design. Manual mode for learning. Never teach bypassPermissions as normal.")

    add_heading_styled(doc, "5.5 Plan mode (deep explanation)", 2)
    p(doc, "Plan mode is one of the highest-leverage features for quality. Claude explores the codebase, asks questions if needed, and writes a plan of attack — without editing your source files until you approve. Enter with /plan or Shift+Tab until plan mode is active.")
    p(doc, "Why it matters for beginners: jumping straight to code often builds the wrong thing quickly. Planning separates “what should we do?” from “do it.” Even the people who built Claude Code publicly emphasize planning first for serious work.")
    p(doc, "Classroom script:")
    code_block(doc, [
        "/plan Build a tip calculator as a single HTML file.",
        "List UI elements, edge cases, and how I will test it.",
        "Do not write code until I approve.",
    ])

    add_heading_styled(doc, "5.6 CLAUDE.md memory", 2)
    p(doc, "CLAUDE.md is a special markdown file Claude reads at the start of every conversation in that project (and optionally from your home folder for global rules). It stores persistent instructions: how to run the project, code style that differs from defaults, gotchas, branch naming rules, etc.")
    p(doc, "Create a starter with /init, then edit it to be short and sharp. Bloated CLAUDE.md files cause Claude to ignore instructions. Rule of thumb for each line: “Would removing this cause repeated mistakes?” If no, delete it.")
    p(doc, "Locations:")
    bullets(doc, [
        "Project: CLAUDE.md or .claude/CLAUDE.md in the repo",
        "User global: ~/.claude/CLAUDE.md (all projects)",
        "Can import other files with @path syntax",
    ])

    add_heading_styled(doc, "5.7 Skills", 2)
    p(doc, "Skills are reusable instruction packages (SKILL.md files) that extend Claude’s capabilities. Unlike CLAUDE.md (always loaded), a skill’s full body loads when relevant or when you invoke /skill-name. Use skills for procedures you repeat: deploy checklist, explain-like-beginner, release notes, etc.")
    p(doc, "Where skills live:")
    bullets(doc, [
        "Personal: ~/.claude/skills/<name>/SKILL.md",
        "Project: .claude/skills/<name>/SKILL.md",
        "Bundled: built-in skills like /doctor, /code-review, /batch, /debug (depending on version)",
    ])

    add_heading_styled(doc, "5.8 Subagents", 2)
    p(doc, "Claude can delegate side work to subagents — separate helper runs with focused context. This matters for large tasks (review in parallel, research while implementing). Beginners: know the concept exists; master single-session building first.")

    add_heading_styled(doc, "5.9 MCP (Model Context Protocol)", 2)
    p(doc, "MCP is a standard way to connect Claude to external tools and data sources (databases, browsers, APIs, company systems). Think of MCP servers as plugins that give Claude new tools. Powerful, but add complexity. In beginner courses, introduce MCP only after students can build and deploy a basic app. When you do: install a trusted MCP server, configure it, use /mcp to manage, and keep secrets out of prompts.")

    add_heading_styled(doc, "5.10 Checkpoints, rewind, resume", 2)
    bullets(doc, [
        "/rewind — roll conversation and/or code back toward a checkpoint when direction went wrong",
        "claude --continue — resume the most recent session in this directory",
        "claude --resume / /resume — pick an older session",
        "/clear — start a fresh conversation (old one still resumable)",
    ])

    add_heading_styled(doc, "5.11 Git integration features", 2)
    p(doc, "Claude can run Git commands conversationally (“commit my changes with a good message”, “create a PR”). It can show diffs (/diff), review code (/code-review), and help resolve conflicts. Students still need to understand Git concepts (Part 14–15) so they know what they are approving.")

    add_heading_styled(doc, "5.12 Images and rich context", 2)
    p(doc, "You can paste screenshots into Claude Code (design mockups, error dialogs). You can @ mention files. You can give documentation URLs. Rich context beats vague descriptions.")
    page_break(doc)

    # ========== PART 6 ==========
    add_heading_styled(doc, "Part 6 — Slash Commands Encyclopedia", 1)
    p(doc, "Slash commands are special instructions that start with /. They control Claude Code itself rather than describing app features. Type / to see available commands in many versions. Below: the teaching set — what it is, why it exists, when to use it, example.")

    def cmd(doc, name, what, why, when, example=None, level="Essential"):
        add_heading_styled(doc, f"{name}  [{level}]", 3)
        multi(doc, [("What: ", True), (what,)])
        multi(doc, [("Why it exists: ", True), (why,)])
        multi(doc, [("When to use: ", True), (when,)])
        if example:
            p(doc, "Example:", bold=True)
            code_block(doc, example)

    cmd(doc, "/help",
        "Shows help for Claude Code commands and usage.",
        "New users get lost; this is the in-product manual entry point.",
        "Any time you forget a command name.",
        "/help", "Essential")

    cmd(doc, "/login",
        "Starts or restarts authentication with your Claude account.",
        "Sessions and machines need credentials; tokens expire or switch accounts.",
        "First install, new computer, auth errors, switching accounts.",
        "/login", "Essential")

    cmd(doc, "/doctor",
        "Runs a setup checkup: installation, configuration, common problems; can help fix them.",
        "Environment issues waste hours; automated diagnosis saves beginners.",
        "Install problems, weird failures, “it worked yesterday.”",
        "/doctor", "Essential")

    cmd(doc, "/init",
        "Analyzes the project and generates a starter CLAUDE.md with detected stack and patterns.",
        "Empty projects lack memory; /init bootstraps useful defaults you then refine.",
        "First time in a new project, or when CLAUDE.md is missing.",
        "/init", "Essential")

    cmd(doc, "/plan [description]",
        "Enters plan mode (optionally with a task description). Claude explores and plans without editing source until approved.",
        "Prevents wrong-direction coding; separates design from implementation.",
        "Any non-trivial feature, refactors, unfamiliar codebases.",
        "/plan Add user login with email and password. Ask clarifying questions first.",
        "Essential")

    cmd(doc, "/clear [name]  (aliases: /reset, /new)",
        "Starts a new conversation with empty context. Previous chat remains available via /resume. Optional name labels the old chat.",
        "Long chats degrade quality; fresh context for a new task is healthier.",
        "Switching tasks; after a messy failed approach; starting a new feature.",
        "/clear auth-feature", "Essential")

    cmd(doc, "/compact [instructions]",
        "Summarizes the conversation to free context space while continuing the same work. Optional focus instructions guide what to keep.",
        "Context fills with file reads and logs; compaction is a middle path between continuing blindly and /clear.",
        "Long debugging session that must continue; “Claude is getting forgetful.”",
        "/compact Keep the decisions about database schema and the failing test name.",
        "Essential")

    cmd(doc, "/permissions  (alias: /allowed-tools)",
        "Opens management of allow/ask/deny rules for tools and related permission settings.",
        "You need control over what Claude may run without asking (network, bash, etc.).",
        "Repeated annoying prompts for safe commands; locking down risky tools.",
        "/permissions", "Important")

    cmd(doc, "/config  (alias: /settings)",
        "Opens settings (theme, model, output style, etc.) or sets keys via key=value in newer versions.",
        "Personal workflow preferences should be adjustable without editing raw JSON always.",
        "Change model, theme, thinking settings, etc.",
        "/config", "Important")

    cmd(doc, "/model",
        "Switch or inspect which model is used for the session (availability depends on plan).",
        "Different models trade speed, cost, and depth.",
        "Hard reasoning tasks vs quick edits.",
        "/model", "Important")

    cmd(doc, "/cost or /usage",
        "Shows usage/cost related information for the session/account context.",
        "AI usage is not infinite; awareness prevents surprise limits.",
        "Long workshops; monitoring spend.",
        "/usage", "Useful")

    cmd(doc, "/resume",
        "Opens a picker (or resumes) previous conversations.",
        "Work spans multiple sittings; continuity matters.",
        "Next day continuation; returning to an old feature branch of conversation.",
        "/resume", "Essential")

    cmd(doc, "/rewind",
        "Rolls back conversation and/or code toward a checkpoint; undoes a bad direction.",
        "AI exploration can dig a hole; rewind is an escape hatch.",
        "Claude made a mess; you want the last good state.",
        "/rewind", "Essential")

    cmd(doc, "/diff",
        "Interactive viewer for uncommitted changes and per-turn diffs.",
        "You must see what actually changed before committing.",
        "Before git commit; reviewing Claude’s last edits.",
        "/diff", "Essential")

    cmd(doc, "/context [all]",
        "Visualizes context usage and may suggest optimizations.",
        "Context is the scarce resource; seeing it teaches better habits.",
        "When quality drops; teaching context management.",
        "/context", "Useful")

    cmd(doc, "/mcp",
        "Manage MCP servers (external tool connections).",
        "MCP extends Claude beyond the repo; needs deliberate setup.",
        "After you intentionally add MCP integrations.",
        "/mcp", "Advanced")

    cmd(doc, "/memory",
        "Work with memory / CLAUDE.md related instructions (refine what Claude remembers).",
        "Persistent memory should be curated, not accidental.",
        "After /init; when Claude keeps forgetting project rules.",
        "/memory", "Important")

    cmd(doc, "/copy [N]",
        "Copies the last assistant response (or Nth latest) to clipboard; can pick code blocks.",
        "Students need snippets in docs or notes without retyping.",
        "Saving explanations; extracting a command block.",
        "/copy", "Useful")

    cmd(doc, "/desktop or /app",
        "Continue the session in the Claude Code Desktop app (Mac/Windows, subscription).",
        "Some prefer GUI diffs and multi-session desktop UX.",
        "When teaching both CLI and desktop surfaces.",
        "/desktop", "Optional")

    cmd(doc, "/code-review …",
        "Skill: review current diff for bugs and cleanups; options to fix or comment on PRs; effort levels.",
        "Second pair of eyes before shipping.",
        "Before merge/demo; after large Claude changes.",
        "/code-review", "Important")

    cmd(doc, "/security-review",
        "Reviews the diff for security vulnerabilities.",
        "AI-generated code often misses authz and injection issues.",
        "Before deploy of apps with users/data.",
        "/security-review", "Important")

    cmd(doc, "/batch <instruction>",
        "Skill: decomposes large multi-file work into units, plans, then can run parallel subagents/worktrees and PRs.",
        "Huge migrations are hard in one chat.",
        "Advanced large refactors — not week 1.",
        "/batch migrate all buttons to new design system", "Advanced")

    cmd(doc, "/debug [description]",
        "Skill: enable debug logging and troubleshoot using session debug logs.",
        "Some failures need deeper instrumentation.",
        "Mysterious Claude Code misbehavior.",
        "/debug Claude keeps failing to run npm", "Useful")

    cmd(doc, "/btw <question>",
        "Ask a side question without polluting the main conversation history the same way.",
        "You want a quick fact without derailing the task.",
        "“btw what does this error code mean?” mid-feature.",
        "/btw what is CORS in one paragraph?", "Useful")

    cmd(doc, "/add-dir <path>",
        "Grant file access to an additional working directory for this session.",
        "Sometimes assets or monorepo packages live outside the main folder.",
        "Multi-folder work; carefully shared resources.",
        "/add-dir ../shared-design", "Advanced")

    cmd(doc, "/cd <path>",
        "Move the session’s working directory (newer versions), preserving more conversation cache carefully.",
        "You started in the wrong folder or need to relocate work.",
        "Project structure shifts mid-session.",
        "/cd ~/projects/my-app", "Useful")

    cmd(doc, "/agents",
        "Create/manage subagent configurations (UI/behavior depends on version).",
        "Specialized agents for review, research, etc.",
        "After basics; team workflows.",
        "/agents", "Advanced")

    cmd(doc, "/tasks",
        "List background tasks running for the session.",
        "Parallel work needs visibility.",
        "When Claude spawned background work.",
        "/tasks", "Advanced")

    cmd(doc, "/background or /bg",
        "Detach session to run as background agent; free the terminal.",
        "Long jobs should not block you.",
        "Long installs/refactors on capable setups.",
        "/background finish running tests and summarize", "Advanced")

    p(doc, "Note: Command availability and exact behavior evolve with Claude Code versions. Teach students to type /help and to ask Claude “what commands are available?” when something is missing. Official reference: code.claude.com/docs/en/commands")
    page_break(doc)

    # ========== PART 7 ==========
    add_heading_styled(doc, "Part 7 — CLAUDE.md, Skills, Permissions, MCP (How to Configure Power)", 1)

    add_heading_styled(doc, "7.1 Writing an excellent CLAUDE.md", 2)
    p(doc, "Good CLAUDE.md example for a beginner project:")
    code_block(doc, [
        "# Project",
        "FocusBoard — personal task app for students learning vibe coding.",
        "Stack: [fill after Claude chooses]  Users: single student user first.",
        "",
        "# Commands",
        "- Install dependencies: ...",
        "- Run dev server: ...",
        "- Run tests: ...",
        "",
        "# Rules",
        "- Prefer simple, readable code over clever code",
        "- After each feature: update README with run steps",
        "- Never commit .env or secrets",
        "- Explain changes in plain English before large refactors",
        "- Wait for approval between milestones",
        "",
        "# Do not",
        "- Do not add payment systems",
        "- Do not add unnecessary libraries",
    ])
    p(doc, "Include: commands Claude cannot guess, project-specific architecture decisions, env var names (not values), test approach.")
    p(doc, "Exclude: tutorials, huge API docs, things Claude can see by reading code, frequent-changing noise.")

    add_heading_styled(doc, "7.2 Skills vs CLAUDE.md vs slash commands", 2)
    table_data(doc, ["Tool", "Loaded when", "Best for"], [
        ["CLAUDE.md", "Every session start", "Stable project facts & rules"],
        ["Skill (SKILL.md)", "When relevant or /name", "Procedures & long checklists"],
        ["Slash command", "When you type /cmd", "Product controls & workflows"],
    ])

    add_heading_styled(doc, "7.3 Example beginner skill: explain-like-beginner", 2)
    p(doc, "Create .claude/skills/explain-like-beginner/SKILL.md with instructions such as: “After any code change, explain in plain language what changed, why, and give 3 manual test steps. Avoid jargon; define terms.” Invoke with /explain-like-beginner or let Claude auto-use when appropriate.")

    add_heading_styled(doc, "7.4 Permissions teaching ladder", 2)
    numbered(doc, [
        "Weeks 1–4: Manual mode only. Students approve edits.",
        "Weeks 5–7: Plan mode before every milestone; Manual for execution.",
        "Week 8+: Optional acceptEdits for speed if student reviews git diff after.",
        "Never: bypassPermissions on personal machines in class.",
    ])
    page_break(doc)

    # ========== PART 8 ==========
    add_heading_styled(doc, "Part 8 — How to Talk to Claude Code (Prompt Craft)", 1)

    add_heading_styled(doc, "8.1 Why vague prompts fail", 2)
    p(doc, "Claude Code will still produce something for “make a nice app,” but it must invent every missing decision: users, features, stack, design, edge cases. Those inventions often do not match what you wanted. Specificity is kindness to the model and to your future self.")

    add_heading_styled(doc, "8.2 The PROMPT framework (expanded)", 2)
    table_data(doc, ["Letter", "Meaning", "Teaching question", "Example fragment"], [
        ["P", "Purpose", "What outcome?", "Tip calculator webpage"],
        ["R", "Rules", "Constraints?", "One HTML file, no frameworks"],
        ["O", "Objects", "What pieces?", "Bill input, tip %, totals"],
        ["M", "Musts", "Acceptance?", "Invalid input shows message"],
        ["P", "Proof", "How verify?", "100 bill + 15% = 115 total"],
        ["T", "Tone", "Style?", "Large fonts, calm colors"],
    ])

    add_heading_styled(doc, "8.3 Patterns that work (from Anthropic best practices)", 2)
    bullets(doc, [
        "Scope the task: file + scenario + constraints.",
        "Point to sources: git history, docs URLs, example files.",
        "Reference existing patterns: “follow HotDogWidget.php style.”",
        "Describe symptoms for bugs: steps, expected, actual, error text.",
        "Ask for evidence: “run tests and paste output,” not “is it done?”",
        "Use @file to attach exact files.",
        "One milestone per message for big work.",
    ])

    add_heading_styled(doc, "8.4 Bad → good library (teach with these)", 2)
    table_data(doc, ["Bad", "Good"], [
        ["fix the app", "Login fails after 30 minutes idle. Check token refresh in auth. Reproduce with a test, then fix root cause."],
        ["make UI better", "Increase spacing, 16px+ body font, primary button contrast. Keep colors. Mobile width 375px must not overflow."],
        ["add database", "Add persistent storage for tasks (title, done, created_at). Explain where data lives. Provide run + migrate steps."],
        ["do auth", "/plan email/password signup+login+logout; users only see own rows. List security concerns. No code until approved."],
    ])

    add_heading_styled(doc, "8.5 Bug report template (force this habit)", 2)
    code_block(doc, [
        "Bug report:",
        "Environment: OS, browser, how I start the app",
        "Steps I took:",
        "Expected:",
        "Actual:",
        "Error text / screenshot notes:",
        "Please find root cause (not a band-aid), fix it, re-run verification, and explain simply.",
    ])
    page_break(doc)

    # ========== PART 9 ==========
    add_heading_styled(doc, "Part 9 — How the Web Works + Full-Stack Mental Model", 1)

    add_heading_styled(doc, "9.1 What happens when you visit a website", 2)
    numbered(doc, [
        "You type a URL or click a link (example.com).",
        "DNS finds the server’s address (like a phone book for computers).",
        "Your browser requests a page (HTTP/HTTPS request).",
        "A server responds with HTML (and references to CSS/JS/images).",
        "Browser renders HTML structure, applies CSS style, runs JavaScript behavior.",
        "Further actions (submit form, load data) may call APIs — more requests to servers.",
    ])

    add_heading_styled(doc, "9.2 Frontend vs backend vs database", 2)
    p(doc, "Restaurant metaphor (useful, but now precise):")
    bullets(doc, [
        "Frontend = dining room: menus, tables, what guests interact with. Technologies: HTML, CSS, JavaScript (and frameworks like React later).",
        "Backend = kitchen: recipes, rules, preparation guests do not see. Technologies: Python, Node.js, etc.",
        "Database = pantry/ledger: lasting ingredients/records. Technologies: SQLite, PostgreSQL, MongoDB, etc.",
        "API = waiter order tickets: structured messages between dining room and kitchen, often JSON over HTTP.",
    ])

    add_heading_styled(doc, "9.3 Client–server in one paragraph", 2)
    p(doc, "Your browser is a client. It requests resources. A server is a program waiting for requests and sending responses. Full-stack means you build both sides (and usually storage). Localhost means the server is running on your own computer during development, often at an address like http://localhost:3000 — the number is a port (a numbered “door” on the machine).")

    add_heading_styled(doc, "9.4 JSON — the lingua franca of APIs", 2)
    p(doc, "JSON is a text format for structured data. Example:")
    code_block(doc, [
        '{',
        '  "title": "Buy milk",',
        '  "done": false,',
        '  "id": 3',
        '}',
    ])
    p(doc, "You do not need to memorize every rule. You need to recognize that apps pass objects/lists around as JSON constantly.")
    page_break(doc)

    # ========== PART 10 HTML ==========
    add_heading_styled(doc, "Part 10 — HTML: The Structure of Every Webpage", 1)

    add_heading_styled(doc, "10.1 What HTML is", 2)
    p(doc, "HTML means HyperText Markup Language. It is not a programming language (it does not compute loops and logic the way JavaScript/Python do). It is a markup language: it labels parts of a document so browsers know what is a heading, a paragraph, a button, an image, a link, a form field, etc.")
    p(doc, "Why you must learn HTML concepts even when Claude writes it:")
    bullets(doc, [
        "You can read the skeleton of any webpage.",
        "You can tell Claude precisely: “the submit button has no type attribute” or “add a label for accessibility.”",
        "You understand SEO basics, accessibility, and structure.",
    ])

    add_heading_styled(doc, "10.2 Tags, elements, attributes", 2)
    p(doc, "An HTML element usually looks like:")
    code_block(doc, '<p class="intro">Hello world</p>')
    bullets(doc, [
        "<p> is the opening tag (paragraph).",
        "</p> is the closing tag.",
        "class=\"intro\" is an attribute (extra information).",
        "Hello world is the content.",
    ])
    p(doc, "Some elements are self-closing / void (no content), like images: <img src=\"photo.jpg\" alt=\"A cat\">")

    add_heading_styled(doc, "10.3 Minimal page skeleton", 2)
    code_block(doc, [
        "<!DOCTYPE html>",
        '<html lang="en">',
        "<head>",
        '  <meta charset="UTF-8" />',
        '  <meta name="viewport" content="width=device-width, initial-scale=1.0" />',
        "  <title>My First Page</title>",
        '  <link rel="stylesheet" href="styles.css" />',
        "</head>",
        "<body>",
        "  <h1>Hello</h1>",
        "  <p>This is my page.</p>",
        '  <script src="app.js"></script>',
        "</body>",
        "</html>",
    ])
    p(doc, "Explanations:")
    bullets(doc, [
        "<!DOCTYPE html> tells the browser this is modern HTML.",
        "<head> contains metadata: title, character set, CSS links — not visible main content.",
        "<body> contains visible content.",
        "viewport meta helps mobile layout.",
        "script at the end often loads JavaScript after HTML exists.",
    ])

    add_heading_styled(doc, "10.4 Essential tags (must know names)", 2)
    table_data(doc, ["Tag", "Purpose", "Notes"], [
        ["h1–h6", "Headings", "One main h1 per page ideally"],
        ["p", "Paragraph", "Body text"],
        ["a", "Link", "href attribute is the destination"],
        ["img", "Image", "src + alt (alt is required for accessibility)"],
        ["button", "Button", "Often needs JS or form handling"],
        ["input", "Form field", "type=text/email/password/number..."],
        ["form", "Form wrapper", "collects inputs to submit"],
        ["label", "Label for input", "Improves usability & a11y"],
        ["ul/ol/li", "Lists", "bulleted / numbered"],
        ["div", "Generic container", "No meaning by itself — overused by beginners"],
        ["header/main/footer/nav/section", "Semantic layout", "Prefer these over endless divs"],
        ["table/tr/td", "Tables", "For tabular data, not page layout"],
    ])

    add_heading_styled(doc, "10.5 Forms — how user data enters apps", 2)
    p(doc, "Forms are how users type data. Critical attributes:")
    bullets(doc, [
        "name — identifies the field when submitting",
        "type — text, email, password, number, checkbox, etc.",
        "required — browser-level “must fill”",
        "placeholder — hint text (not a substitute for a label)",
    ])
    code_block(doc, [
        "<form>",
        '  <label for="email">Email</label>',
        '  <input id="email" name="email" type="email" required />',
        '  <button type="submit">Sign up</button>',
        "</form>",
    ])

    add_heading_styled(doc, "10.6 What to ask Claude regarding HTML", 2)
    bullets(doc, [
        "Use semantic tags for accessibility",
        "Every image needs meaningful alt text",
        "Associate labels with inputs",
        "Keep structure simple and readable",
        "Explain the HTML outline of the page to me",
    ])
    page_break(doc)

    # ========== PART 11 CSS ==========
    add_heading_styled(doc, "Part 11 — CSS: Making Pages Look Right", 1)

    add_heading_styled(doc, "11.1 What CSS is", 2)
    p(doc, "CSS means Cascading Style Sheets. It describes presentation: colors, fonts, spacing, layout, responsive behavior. HTML is structure; CSS is appearance.")

    add_heading_styled(doc, "11.2 How CSS attaches to HTML", 2)
    bullets(doc, [
        "External file: styles.css linked with <link> (best for real projects)",
        "Internal <style> block in head (fine for tiny demos)",
        "Inline style=\"...\" on an element (avoid for large work)",
    ])

    add_heading_styled(doc, "11.3 Selectors and declarations", 2)
    code_block(doc, [
        "p {",
        "  color: #1a2421;",
        "  font-size: 16px;",
        "  line-height: 1.5;",
        "}",
        "",
        ".card {",
        "  padding: 16px;",
        "  border-radius: 8px;",
        "}",
        "",
        "#main-title {",
        "  font-size: 32px;",
        "}",
    ])
    bullets(doc, [
        "p selects all paragraphs",
        ".card selects elements with class=\"card\"",
        "#main-title selects id=\"main-title\" (ids should be unique)",
    ])

    add_heading_styled(doc, "11.4 The box model (critical mental model)", 2)
    p(doc, "Every element is a box with content, padding (inner space), border, and margin (outer space). Beginner layout bugs are often padding/margin confusion. When something looks “too tight” or “won’t center,” think box model.")

    add_heading_styled(doc, "11.5 Layout tools you will hear Claude use", 2)
    table_data(doc, ["Tool", "What it’s for", "Beginner note"], [
        ["Flexbox", "One-dimensional row/column layouts", "Great for nav bars, centering"],
        ["Grid", "Two-dimensional layouts", "Great for page sections"],
        ["Media queries", "Different CSS at different screen widths", "Mobile responsiveness"],
        ["position", "Relative/absolute/fixed placement", "Easy to misuse — be careful"],
    ])

    add_heading_styled(doc, "11.6 Design quality bar for non-designers", 2)
    bullets(doc, [
        "Body text at least ~16px",
        "Strong contrast (dark text on light, or careful dark theme)",
        "Consistent spacing scale (8px rhythm is a common habit)",
        "Primary button looks clickable and obvious",
        "On mobile, no horizontal scrolling of the whole page",
        "Error messages in plain language near the field",
    ])
    page_break(doc)

    # ========== PART 12 JS ==========
    add_heading_styled(doc, "Part 12 — JavaScript: Making Pages Do Things", 1)

    add_heading_styled(doc, "12.1 What JavaScript is", 2)
    p(doc, "JavaScript (JS) is a programming language that runs in the browser (and also on servers via Node.js). It handles behavior: what happens when you click, type, submit, fetch data from an API, update text on the page without a full reload, etc.")
    p(doc, "HTML = nouns (things on the page). CSS = adjectives (how they look). JavaScript = verbs (what they do).")

    add_heading_styled(doc, "12.2 Core concepts (plain language)", 2)
    table_data(doc, ["Concept", "Meaning", "Tiny example idea"], [
        ["Variable", "A named box holding a value", "let score = 0"],
        ["String", "Text data", "\"hello\""],
        ["Number", "Numeric data", "42"],
        ["Boolean", "true/false", "isLoggedIn"],
        ["Array", "Ordered list", "[1, 2, 3]"],
        ["Object", "Named fields bag", "{ title: \"Task\" }"],
        ["Function", "Reusable procedure", "function add(a,b)"],
        ["if/else", "Decisions", "if score > 10 ..."],
        ["loop", "Repeat work", "for each item in list"],
        ["DOM", "JS view of the HTML page", "document.querySelector"],
        ["event", "Something that happened", "click, submit, keydown"],
        ["async/await", "Wait for slow work (network)", "fetch data from API"],
    ])

    add_heading_styled(doc, "12.3 DOM and events — the heart of interactive pages", 2)
    p(doc, "The DOM (Document Object Model) is the browser’s live representation of your HTML. JavaScript can find elements, change text, show/hide sections, and listen for events.")
    code_block(doc, [
        "const button = document.querySelector(\"#save\");",
        "button.addEventListener(\"click\", () => {",
        "  alert(\"Saved!\");",
        "});",
    ])
    p(doc, "You do not need to write this from scratch on day one. You need to understand: button click → function runs → page or data changes.")

    add_heading_styled(doc, "12.4 Fetching data (how frontends talk to backends)", 2)
    p(doc, "Modern apps often call APIs using fetch. The browser requests JSON from a server URL, then updates the page. When students say “the list is empty but the server has data,” the bug is often in this bridge.")

    add_heading_styled(doc, "12.5 JavaScript vs Java", 2)
    p(doc, "They are different languages. The similar names are historical accidents. This course means JavaScript when we say JS.")

    add_heading_styled(doc, "12.6 Node.js and npm (just enough)", 2)
    p(doc, "Node.js lets you run JavaScript outside the browser (servers, tooling). npm is a package manager: it installs libraries other people wrote into your project (package.json lists them). Commands you will see:")
    code_block(doc, [
        "node -v",
        "npm -v",
        "npm install",
        "npm run dev",
    ])
    p(doc, "Claude will often run these for you. You should know what “installing dependencies” means: downloading the libraries your project needs.")
    page_break(doc)

    # ========== PART 13 PYTHON ==========
    add_heading_styled(doc, "Part 13 — Python: Backend Language Foundations", 1)

    add_heading_styled(doc, "13.1 Why Python appears in this course", 2)
    p(doc, "Python is a popular, readable language used for backends, automation, data, and AI tooling. Many beginner-friendly servers are written in Python (Flask, FastAPI). Even if your cohort chooses a Node stack, Python literacy helps you read examples and scripts.")

    add_heading_styled(doc, "13.2 What a Python program looks like", 2)
    code_block(doc, [
        "# greet.py",
        "def greet(name):",
        "    return f\"Hello, {name}!\"",
        "",
        "print(greet(\"Asha\"))",
    ])
    bullets(doc, [
        "def defines a function",
        "Indentation (spaces) is syntax in Python — it defines blocks",
        "print shows output in the terminal",
        "Run with: python greet.py  (or python3 greet.py)",
    ])

    add_heading_styled(doc, "13.3 Core concepts map", 2)
    table_data(doc, ["Concept", "Python flavor", "Why it matters in apps"], [
        ["Variables", "name = \"Ada\"", "Store user input and state"],
        ["Lists", "[1, 2, 3]", "Collections of records"],
        ["Dicts", "{\"title\": \"Task\"}", "JSON-like objects"],
        ["Functions", "def save():", "Reusable backend logic"],
        ["Modules", "import os", "Reuse libraries"],
        ["Virtual env", "venv", "Isolate project dependencies"],
        ["pip", "pip install flask", "Install packages"],
        ["requirements.txt", "pinned deps list", "Others can recreate environment"],
    ])

    add_heading_styled(doc, "13.4 Python as a tiny web backend (concept)", 2)
    p(doc, "A backend framework receives HTTP requests and returns responses. Conceptual flow:")
    numbered(doc, [
        "Browser sends GET /api/tasks",
        "Python server function runs",
        "It reads from database",
        "It returns JSON list of tasks",
        "Frontend JavaScript renders the list",
    ])
    p(doc, "You can ask Claude: “Explain each file in this Flask/FastAPI app like I am new. What is a route? What is a request body?”")

    add_heading_styled(doc, "13.5 When to choose Python vs JavaScript backend", 2)
    table_data(doc, ["Choose more JS/Node when", "Choose more Python when"], [
        ["Cohort already uses JS on frontend", "Instructor is strongest in Python"],
        ["You want one language ecosystem", "You want very readable server examples"],
        ["Many modern full-stack JS templates", "Data/AI adjacent projects later"],
    ])
    p(doc, "Teaching rule: pick ONE primary stack for the whole class to reduce support chaos.")
    page_break(doc)

    # ========== PART 14 GIT ==========
    add_heading_styled(doc, "Part 14 — Git: Version Control Explained Properly", 1)

    add_heading_styled(doc, "14.1 The problem Git solves", 2)
    p(doc, "Without Git, people save final2_really_final.zip forever and lose history. Git records snapshots of your project over time so you can:")
    bullets(doc, [
        "Go back to a previous working version",
        "See what changed and who changed it",
        "Work on features without destroying the main version",
        "Collaborate without overwriting each other blindly",
    ])

    add_heading_styled(doc, "14.2 Core vocabulary", 2)
    table_data(doc, ["Term", "Plain meaning", "Analogy"], [
        ["Repository (repo)", "A project tracked by Git", "A binder with history"],
        ["Working tree", "Your current files", "Papers on the desk"],
        ["Staging area", "Files marked for next snapshot", "Papers in the “to file” tray"],
        ["Commit", "A saved snapshot + message", "A labeled save point"],
        ["Branch", "A movable line of development", "Parallel timeline"],
        ["main/master", "Default primary branch", "Official timeline"],
        ["Merge", "Combine histories", "Bring side work into main"],
        ["Clone", "Copy a repo to your machine", "Download the binder"],
        ["Remote", "A hosted copy (often GitHub)", "Cloud backup"],
        ["Push", "Upload commits to remote", "Sync up"],
        ["Pull", "Download new commits", "Sync down"],
        ["Conflict", "Two edits clash", "Two people edited same paragraph"],
    ])

    add_heading_styled(doc, "14.3 The basic daily loop", 2)
    code_block(doc, [
        "git status                 # what changed?",
        "git add .                  # stage changes (or add specific files)",
        "git commit -m \"Add login form validation\"",
        "git push                   # upload to GitHub (after remote set)",
    ])
    p(doc, "You can do this in English with Claude Code: “Show me what changed. Commit with a clear message. Do not push yet.” But you must understand the meaning of each step.")

    add_heading_styled(doc, "14.4 Commit messages that help humans", 2)
    p(doc, "Bad: \"fix\", \"updates\", \"asdf\". Good: describe why/what outcome: \"Prevent empty tasks from being saved\". Future you will thank present you.")

    add_heading_styled(doc, "14.5 .gitignore — critical for security", 2)
    p(doc, "Git tracks files you add. Some files must never be committed:")
    bullets(doc, [
        ".env and secret keys",
        "node_modules/ (reinstallable dependencies)",
        "Python venv folders",
        "OS junk like .DS_Store sometimes",
        "Large personal datasets",
    ])
    p(doc, "A .gitignore file lists patterns to skip. On day one of GitHub, verify secrets are not tracked: ask Claude “Is .env ignored? Show git status.”")

    add_heading_styled(doc, "14.6 Branching (simple teaching model)", 2)
    p(doc, "main stays stable. You create a branch feature/login, do work, then merge back (or open a Pull Request on GitHub). Beginners can work only on main for tiny solo projects, but should learn branch vocabulary before team work.")

    add_heading_styled(doc, "14.7 Common Git errors and meanings", 2)
    table_data(doc, ["Message / situation", "Meaning", "What to do"], [
        ["not a git repository", "You did not git init / wrong folder", "cd to project; git init or clone"],
        ["Your branch is ahead", "Local commits not pushed", "git push when ready"],
        ["merge conflict", "Same lines changed two ways", "Edit conflict markers; ask Claude carefully"],
        ["Permission denied (publickey)", "GitHub auth not set", "HTTPS + credential helper or SSH keys"],
        ["large file rejected", "GitHub file size limits", "Remove file; use Git LFS only if needed"],
    ])
    page_break(doc)

    # ========== PART 15 GITHUB ==========
    add_heading_styled(doc, "Part 15 — GitHub: Complete Beginner Guide", 1)

    add_heading_styled(doc, "15.1 What GitHub is (and is not)", 2)
    p(doc, "Git is the version control system (software on your machine and in repos). GitHub is a website/company platform that hosts Git repositories online, adds collaboration features (Pull Requests, Issues, Actions), and acts as a portfolio for builders. Alternatives exist (GitLab, Bitbucket), but GitHub is the industry default for many courses.")

    add_heading_styled(doc, "15.2 Create an account the right way", 2)
    numbered(doc, [
        "Go to github.com and sign up with a professional username (this may appear on job applications).",
        "Verify email.",
        "Enable 2FA (Settings → Password and authentication).",
        "Do not reuse passwords.",
    ])

    add_heading_styled(doc, "15.3 Create a repository", 2)
    numbered(doc, [
        "Click New repository.",
        "Choose a clear name: focusboard or vibe-tip-calculator.",
        "Add a README (yes for beginners).",
        "Add .gitignore template (Node or Python depending on stack).",
        "Choose public (portfolio) or private (class-only) intentionally.",
        "Do not upload secrets.",
    ])

    add_heading_styled(doc, "15.4 Connect local project to GitHub", 2)
    p(doc, "Typical flow (Claude can guide live):")
    code_block(doc, [
        "git init",
        "git add .",
        "git commit -m \"Initial commit\"",
        "git branch -M main",
        "git remote add origin https://github.com/USER/REPO.git",
        "git push -u origin main",
    ])
    p(doc, "Authentication: GitHub no longer accepts account password for Git over HTTPS in the old way; use GitHub CLI (gh auth login), personal access token, or SSH keys. In class, gh auth login is often easiest.")

    add_heading_styled(doc, "15.5 README that looks professional", 2)
    p(doc, "A strong README includes:")
    bullets(doc, [
        "Project title and one-sentence description",
        "Screenshot or GIF",
        "Features list",
        "Tech stack",
        "Setup steps (install, env vars, run)",
        "How to test",
        "Live demo URL",
        "Author name",
        "License (optional but good)",
    ])

    add_heading_styled(doc, "15.6 Issues, Pull Requests, and Code Review (concepts)", 2)
    bullets(doc, [
        "Issue: a tracked bug or task discussion.",
        "Pull Request (PR): a proposal to merge a branch’s changes into another branch, with discussion and review.",
        "Review: humans (or Claude /code-review) comment on the diff before merge.",
        "For solo beginners: PRs still teach professional workflow even if you merge your own.",
    ])

    add_heading_styled(doc, "15.7 GitHub and Claude Code together", 2)
    p(doc, "Useful prompts:")
    code_block(doc, [
        "Create a clean README for beginners.",
        "Prepare .gitignore for this stack and ensure .env is ignored.",
        "Commit current work with a descriptive message.",
        "Help me open a pull request with a good summary.",
        "/code-review",
        "/security-review",
    ])

    add_heading_styled(doc, "15.8 Portfolio habits", 2)
    bullets(doc, [
        "Pin 1–3 best repos on your GitHub profile",
        "Write READMEs as if a stranger will run the project tomorrow",
        "Small frequent commits beat giant mystery dumps",
        "Never commit API keys — recruiters and attackers both look",
    ])
    page_break(doc)

    # ========== PART 16 BUILD PATH ==========
    add_heading_styled(doc, "Part 16 — Build Path: From Hello World to Full-Stack", 1)
    p(doc, "This is the practice spine of the course. Each project adds one layer of complexity. Do not skip early projects; they create intuition.")

    add_heading_styled(doc, "Project 0 — Environment proves itself", 2)
    bullets(doc, [
        "Install Claude Code; claude doctor clean enough to work",
        "Create course folder structure",
        "Generate a Hello HTML page via Claude and open in browser",
    ])

    add_heading_styled(doc, "Project 1 — Multi-page static site", 2)
    p(doc, "Skills: HTML structure, CSS basics, file linking, relative paths.")
    p(doc, "Deliverable: portfolio with Home/About/Projects/Contact.")

    add_heading_styled(doc, "Project 2 — Interactive single-page tool", 2)
    p(doc, "Skills: JS events, input validation, DOM updates.")
    p(doc, "Deliverable: tip calculator or quiz with score.")

    add_heading_styled(doc, "Project 3 — Persistent local app", 2)
    p(doc, "Skills: CRUD mental model, localStorage or simple file/db, state.")
    p(doc, "Deliverable: notes/todo that survives refresh.")

    add_heading_styled(doc, "Project 4 — Backend API + frontend", 2)
    p(doc, "Skills: routes, JSON, separation of concerns, run two layers.")
    p(doc, "Deliverable: tasks API + UI that calls it.")

    add_heading_styled(doc, "Project 5 — Auth + ownership", 2)
    p(doc, "Skills: signup/login, password hashing (via libraries), authorization checks.")
    p(doc, "Deliverable: users only see their own records; test with two accounts.")

    add_heading_styled(doc, "Project 6 — Capstone full-stack product", 2)
    p(doc, "Skills: PRD, milestones, CLAUDE.md, GitHub, deploy, demo.")
    p(doc, "Deliverable: public URL + repo + 5-minute presentation.")

    add_heading_styled(doc, "Milestone discipline (non-negotiable)", 2)
    p(doc, "For every non-trivial feature:")
    numbered(doc, [
        "/plan and approve",
        "Implement only that milestone",
        "Run/verify yourself",
        "Commit",
        "Then next milestone",
    ])
    page_break(doc)

    # ========== PART 17 DEPLOY ==========
    add_heading_styled(doc, "Part 17 — Deploying Your App", 1)

    add_heading_styled(doc, "17.1 What deployment means", 2)
    p(doc, "Deployment is publishing your application so other people can use it over the internet. Your laptop is for development. A host (Vercel, Netlify, Render, Railway, Fly.io, cloud VMs, etc.) provides always-on computers and networking.")

    add_heading_styled(doc, "17.2 What usually must be configured", 2)
    bullets(doc, [
        "Build command (how to compile/bundle the app)",
        "Start command (how to run the server)",
        "Environment variables (secrets and config for production)",
        "Database hosting (if not SQLite-on-disk alone)",
        "Domain / URL",
    ])

    add_heading_styled(doc, "17.3 Why “works on my machine” fails online", 2)
    bullets(doc, [
        "Forgot to set env vars on the host",
        "Database URL points to localhost",
        "Hard-coded file paths",
        "Dev-only dependencies missing in production",
        "CORS misconfiguration",
        "Wrong Node/Python version",
    ])

    add_heading_styled(doc, "17.4 Smoke test after deploy", 2)
    numbered(doc, [
        "Open URL in incognito",
        "Sign up / create data",
        "Refresh — data still there",
        "Test on a phone",
        "Second account isolation (if auth exists)",
        "Put URL in README",
    ])
    page_break(doc)

    # ========== PART 18 ==========
    add_heading_styled(doc, "Part 18 — Capstone, Assessment, Teaching Notes", 1)

    add_heading_styled(doc, "18.1 Capstone requirements", 2)
    bullets(doc, [
        "Real problem statement and user",
        "PRD with MVP scope",
        "Working full-stack flow (CRUD + persistence)",
        "GitHub repo with quality README",
        "Deployed live URL",
        "Demo: problem, walkthrough, 3 human decisions, next steps",
        "Optional: authentication",
    ])

    add_heading_styled(doc, "18.2 Suggested rubric (100)", 2)
    table_data(doc, ["Criterion", "Points"], [
        ["Problem clarity & PRD", "15"],
        ["Working core user flow", "25"],
        ["Data persistence", "15"],
        ["Auth/ownership or honest single-user scope", "10"],
        ["UX readability", "10"],
        ["Git history + README", "10"],
        ["Deployment", "10"],
        ["Demo & reflection quality", "5"],
    ])

    add_heading_styled(doc, "18.3 Instructor principles that raise quality", 2)
    bullets(doc, [
        "Explain the concept, then demo, then lab — never lab without model.",
        "Force students to read errors out loud.",
        "Ban “make it good” prompts; require PROMPT structure.",
        "Use plan mode as a cultural norm.",
        "One stack for the cohort.",
        "TA support during install week is mandatory.",
        "Celebrate deployed links; identity shift matters.",
    ])
    page_break(doc)

    # ========== APPENDICES ==========
    add_heading_styled(doc, "Appendix A — Master Glossary", 1)
    table_data(doc, ["Term", "Definition"], [
        ["Agentic", "AI that takes multi-step actions with tools, not only chat replies"],
        ["API", "Interface for programs to communicate (often HTTP + JSON)"],
        ["Backend", "Server-side logic and data processing"],
        ["Branch", "Parallel line of Git history"],
        ["Bug", "Behavior that does not match intent"],
        ["CLAUDE.md", "Persistent project instructions for Claude Code"],
        ["Commit", "Git snapshot with message"],
        ["Context window", "How much info the model can consider at once"],
        ["CRUD", "Create, Read, Update, Delete"],
        ["CSS", "Language for styling web pages"],
        ["Deploy", "Publish app for others online"],
        ["DOM", "Browser’s live object model of a page"],
        ["Frontend", "User-facing interface"],
        ["Full-stack", "Frontend + backend + data"],
        ["Git", "Version control system"],
        ["GitHub", "Hosted Git + collaboration platform"],
        ["HTML", "Markup language for page structure"],
        ["HTTP", "Protocol browsers use to request resources"],
        ["JavaScript", "Language for web behavior (and Node backends)"],
        ["JSON", "Common text data format"],
        ["Localhost", "Your computer as the server in development"],
        ["MCP", "Model Context Protocol — tool/plugin connections"],
        ["MVP", "Minimum Viable Product"],
        ["npm", "Node package manager"],
        ["PATH", "Where the shell looks for programs"],
        ["Plan mode", "Claude plans without editing until approved"],
        ["Port", "Numbered network endpoint on a machine"],
        ["PR", "Pull Request — proposal to merge changes"],
        ["PRD", "Product Requirements Document"],
        ["Python", "Popular general-purpose language, common for backends"],
        ["Repo", "Repository — project under Git"],
        ["Skill", "On-demand instruction package for Claude"],
        ["Terminal", "Text interface to run commands"],
        ["Token (AI)", "Chunk of text the model processes (also: auth token sense)"],
    ])

    add_heading_styled(doc, "Appendix B — Terminal Cheatsheet", 1)
    code_block(doc, [
        "pwd / cd / ls / mkdir / clear",
        "cat file   OR   type file   OR   Get-Content file",
        "claude",
        "claude --version",
        "claude doctor",
        "claude --continue",
        "git status / git add . / git commit -m \"msg\" / git push",
    ])

    add_heading_styled(doc, "Appendix C — Claude Code Essential Commands Cheatsheet", 1)
    code_block(doc, [
        "/help  /login  /doctor  /init  /plan  /clear  /compact",
        "/diff  /rewind  /resume  /permissions  /config  /usage",
        "/memory  /mcp  /code-review  /security-review  /context",
    ])

    add_heading_styled(doc, "Appendix D — Install & first-run troubleshooting", 1)
    table_data(doc, ["Problem", "Likely fix"], [
        ["command not found: claude", "Add install dir to PATH; new terminal; reinstall"],
        ["irm not recognized", "You are in CMD — use PowerShell or CMD installer"],
        ["&& not valid in PowerShell", "You mixed CMD syntax into PowerShell"],
        ["Login / plan access errors", "Need Pro/Max/Team/Console — free insufficient"],
        ["Git Bash missing on Windows", "Install Git for Windows; set path if needed"],
        ["Permission prompts forever", "Normal in Manual mode — or carefully tune /permissions"],
    ])

    add_heading_styled(doc, "Appendix E — Recommended 10-week teaching sequence", 1)
    table_data(doc, ["Week", "Teaching focus", "Student proof"], [
        ["0", "Accounts + machine readiness", "Checklist complete"],
        ["1", "Files, terminal, Claude install, HTML hello", "Hello page"],
        ["2", "Slash commands + plan mode + PROMPT", "Prompt journal + tool"],
        ["3", "Web model + HTML/CSS deeper + portfolio", "Multi-page site"],
        ["4", "JavaScript interactivity", "Quiz/calculator"],
        ["5", "PRD + Git foundations", "PRD + first commits"],
        ["6", "Python/Node backend concepts + persistence", "CRUD app"],
        ["7", "Full-stack MVP milestones", "Local MVP video"],
        ["8", "Auth + CLAUDE.md + skills intro", "Ownership tests"],
        ["9", "GitHub portfolio quality", "Public repo"],
        ["10", "Deploy + security review + demo day", "Live URL + talk"],
    ])

    doc.add_paragraph()
    callout(doc, "Final word to the instructor", [
        "You are teaching people to command a powerful agentic engineer and to understand the world that engineer operates in.",
        "Philosophy without mechanisms produces confidence without competence.",
        "Mechanisms without empathy produce fear.",
        "This guide aims at both: deep explanation and human pacing.",
        "Go slow in weeks 1–2. Demand verification always. Ship something real.",
    ])

    page_break(doc)

    # ========== EXPANDED DEEP DIVES ==========
    add_heading_styled(doc, "Part 19 — Deep Dive: A Complete Claude Code Session Walkthrough", 1)
    p(doc, "This part is a “film study” of a good session. Teach it live by projecting your terminal.")
    add_heading_styled(doc, "19.1 Scenario", 2)
    p(doc, "Build a “Study Timer” page: user enters minutes, presses Start, sees countdown, hears/sees completion message, can reset. Single HTML file is fine for v1.")
    add_heading_styled(doc, "19.2 Session transcript (what you type and why)", 2)
    numbered(doc, [
        "cd into empty folder; run claude. Why: correct working directory.",
        "/init then trim CLAUDE.md to 10–15 lines. Why: persistent run rules.",
        "/plan Create a study timer... list edge cases (0 minutes, non-numbers, tab background). No code yet. Why: surface assumptions.",
        "Reply: “Approve with change: also disable Start while running.” Why: human product decision.",
        "Implement approved plan only. After done, tell me exact file path and how to open. Why: verification ownership.",
        "You open browser; try edge cases; paste bug report if needed. Why: reality check.",
        "/diff. Why: see what changed before any commit.",
        "Commit with message “Add study timer with validation and running-state lock.” Why: history.",
        "/clear when moving to a different feature tomorrow. Why: context hygiene.",
    ])
    add_heading_styled(doc, "19.3 What “good steering” looks like mid-session", 2)
    bullets(doc, [
        "“Stop. Do not refactor unrelated files.”",
        "“Explain the timer logic in plain English before more edits.”",
        "“Only change CSS for mobile; leave JS alone.”",
        "“You introduced a new library — remove it; stay vanilla.”",
        "“Show me the failing case as a manual test checklist.”",
    ])
    p(doc, "Steering is a skill. Silent acceptance trains you to be a passenger.")

    add_heading_styled(doc, "Part 20 — Deep Dive: HTML Forms, Accessibility, and Multi-page Sites", 1)
    add_heading_styled(doc, "20.1 Forms that do not lie to users", 2)
    p(doc, "A form is a contract: the user provides data; the system responds. Browser validation (required, type=email) is a first line of defense, not the only line. Backend validation must repeat checks. For pure frontend tools, validation still prevents nonsense states.")
    code_block(doc, [
        "<form id=\"signup\" novalidate>",
        "  <!-- novalidate only if you intentionally replace native validation with custom JS -->",
        '  <div class="field">',
        '    <label for="email">Email address</label>',
        '    <input id="email" name="email" type="email" autocomplete="email" required />',
        '    <p class="error" id="email-error" hidden>Enter a valid email.</p>',
        "  </div>",
        '  <button type="submit">Create account</button>',
        "</form>",
    ])
    bullets(doc, [
        "label for= + input id= connects text to control (screen readers + click-to-focus).",
        "autocomplete helps browsers fill known fields securely.",
        "error region should be announced; ask Claude for aria-live if building advanced a11y.",
        "type=submit vs type=button: submit triggers form submit event; button does not by default.",
    ])
    add_heading_styled(doc, "20.2 Multi-page navigation model", 2)
    p(doc, "Each HTML file is a page. Shared CSS via <link href=\"styles.css\">. Shared JS via <script src=\"app.js\">. Navigation uses <a href=\"about.html\">. If you deploy under a subpath, relative links are safer than hard-coded domains during learning.")
    add_heading_styled(doc, "20.3 Accessibility non-negotiables for this course", 2)
    numbered(doc, [
        "Images that convey meaning need alt text; decorative images can use alt=\"\".",
        "Buttons should be <button>, not clickable divs.",
        "Color is not the only error signal — also use text.",
        "Focus outlines should remain visible for keyboard users.",
        "Page <title> should be unique and descriptive.",
    ])

    add_heading_styled(doc, "Part 21 — Deep Dive: CSS Layout Studio", 1)
    add_heading_styled(doc, "21.1 Centering — the classic beginner pain", 2)
    code_block(doc, [
        ".page {",
        "  min-height: 100vh;",
        "  display: flex;",
        "  justify-content: center; /* horizontal in a row flex */",
        "  align-items: center;     /* vertical in a row flex */",
        "}",
    ])
    p(doc, "Teach students to name the axis. In a default row flex container, main axis is horizontal; cross axis is vertical. justify-content acts on main; align-items on cross.")
    add_heading_styled(doc, "21.2 A simple responsive grid for cards", 2)
    code_block(doc, [
        ".grid {",
        "  display: grid;",
        "  grid-template-columns: repeat(auto-fit, minmax(240px, 1fr));",
        "  gap: 16px;",
        "}",
    ])
    p(doc, "auto-fit + minmax creates as many columns as fit, each at least 240px. This single pattern upgrades portfolio project cards dramatically.")
    add_heading_styled(doc, "21.3 Specificity (why “my CSS didn’t apply”)", 2)
    p(doc, "When two rules conflict, the more specific wins (roughly: id > class > element), then order. !important is a last resort and creates messes. Prefer clearer structure over !important wars. If Claude’s styles “don’t show,” check: wrong file linked? cached CSS? selector mismatch? more specific override?")

    add_heading_styled(doc, "Part 22 — Deep Dive: JavaScript State & DOM Studio", 1)
    add_heading_styled(doc, "22.1 A minimal mental model of an interactive app", 2)
    code_block(doc, [
        "// 1) state",
        "let count = 0;",
        "",
        "// 2) references to DOM nodes",
        "const out = document.querySelector(\"#out\");",
        "const btn = document.querySelector(\"#inc\");",
        "",
        "// 3) render function: state -> UI",
        "function render() {",
        "  out.textContent = String(count);",
        "}",
        "",
        "// 4) events update state, then render",
        "btn.addEventListener(\"click\", () => {",
        "  count += 1;",
        "  render();",
        "});",
        "",
        "render();",
    ])
    p(doc, "Almost every UI bug is either wrong state, wrong render, or event not wired. When students debug with Claude, force them to identify which of the three failed.")
    add_heading_styled(doc, "22.2 Arrays of objects — how real apps store lists", 2)
    code_block(doc, [
        "const tasks = [",
        "  { id: 1, title: \"Read HTML chapter\", done: false },",
        "  { id: 2, title: \"Ship portfolio\", done: true },",
        "];",
    ])
    p(doc, "Rendering a list means looping tasks and creating DOM nodes or HTML strings. Updating means changing the object, then re-rendering. Deleting means filtering the array. This is CRUD on the client before a backend exists.")
    add_heading_styled(doc, "22.3 Talking to APIs with fetch", 2)
    code_block(doc, [
        "const res = await fetch(\"/api/tasks\");",
        "if (!res.ok) throw new Error(\"Failed to load tasks\");",
        "const data = await res.json();",
        "// data is now a JS array/object from JSON",
    ])
    p(doc, "await pauses the function until the network responds. Errors happen: network down, 500 from server, invalid JSON. Teaching point: always check res.ok and show a user-friendly message.")

    add_heading_styled(doc, "Part 23 — Deep Dive: Python Backend Studio", 1)
    add_heading_styled(doc, "23.1 From script to server", 2)
    p(doc, "A Python script runs top-to-bottom and exits. A server process starts and waits in a loop for HTTP requests until you stop it. Frameworks map URL paths + methods to functions.")
    add_heading_styled(doc, "23.2 Example route logic (conceptual FastAPI style)", 2)
    code_block(doc, [
        "# Conceptual — Claude will generate exact syntax for your chosen framework",
        "@app.get(\"/api/tasks\")",
        "def list_tasks():",
        "    return tasks  # framework converts to JSON",
        "",
        "@app.post(\"/api/tasks\")",
        "def create_task(payload):",
        "    if not payload.title.strip():",
        "        # return 400 bad request",
        "        ...",
        "    # save and return created task with id",
    ])
    add_heading_styled(doc, "23.3 Environment & dependencies", 2)
    code_block(doc, [
        "python3 -m venv .venv",
        "source .venv/bin/activate   # Windows: .venv\\Scripts\\activate",
        "pip install -r requirements.txt",
    ])
    p(doc, "Explain: venv isolates packages; requirements.txt freezes the list so classmates install the same versions. Never commit the venv folder; commit requirements.txt.")
    add_heading_styled(doc, "23.4 Status codes students should recognize", 2)
    table_data(doc, ["Code", "Meaning", "Typical cause"], [
        ["200", "OK", "Successful read"],
        ["201", "Created", "Successful create"],
        ["400", "Bad request", "Invalid input"],
        ["401", "Unauthorized", "Not logged in"],
        ["403", "Forbidden", "Logged in but not allowed"],
        ["404", "Not found", "Wrong id/path"],
        ["500", "Server error", "Bug/exception on server"],
    ])

    add_heading_styled(doc, "Part 24 — Deep Dive: Git Internals for Beginners (Without Fear)", 1)
    p(doc, "You do not need to memorize plumbing commands. You do need this model:")
    bullets(doc, [
        "Working directory: files you edit",
        "Staging index: proposed next commit contents",
        "Commit objects: immutable snapshots linked into history",
        "Branches: named pointers to commits",
        "Remotes: named URLs to other copies of the repo",
    ])
    add_heading_styled(doc, "24.1 A story of three commits", 2)
    code_block(doc, [
        "A---B---C          main",
        "     \\",
        "      D---E        feature/login",
    ])
    p(doc, "main points at C. feature/login points at E. Merging feature into main combines histories. Conflicts happen when B→C and B→D changed the same lines differently.")
    add_heading_styled(doc, "24.2 What Claude should not be allowed to do casually", 2)
    bullets(doc, [
        "Force push to main",
        "git reset --hard without explicit human understanding",
        "Rewrite published history in a shared class repo",
        "Commit .env “just this once”",
    ])

    add_heading_styled(doc, "Part 25 — Deep Dive: GitHub Collaboration & Portfolio Operations", 1)
    add_heading_styled(doc, "25.1 Issues as a task system", 2)
    p(doc, "An Issue is a tracked conversation: bug, feature idea, or chore. Good issue template: summary, steps to reproduce, expected/actual, screenshots. Students can manage capstone work as issues even solo.")
    add_heading_styled(doc, "25.2 Pull Request quality bar", 2)
    bullets(doc, [
        "Title states outcome: “Add task validation errors”",
        "Body explains why, not only what",
        "Screenshots for UI changes",
        "Test plan checklist",
        "Small PRs merge easier than monster PRs",
    ])
    add_heading_styled(doc, "25.3 Branch protection (team mode)", 2)
    p(doc, "On real teams, main may require PR reviews before merge. Even if your class does not enable branch protection, teach that professional defaults exist so students are not shocked at internships.")
    add_heading_styled(doc, "25.4 GitHub Pages vs app hosts", 2)
    p(doc, "GitHub Pages hosts static sites (HTML/CSS/JS). Apps with servers/databases need hosts that run processes (Render, Railway, Fly, etc.) or serverless platforms. Match host to architecture — a common student failure is trying to put a Python API on static-only hosting.")

    add_heading_styled(doc, "Part 26 — Deep Dive: Debugging Curriculum", 1)
    p(doc, "Debugging is not panic. It is a loop:")
    numbered(doc, [
        "Reproduce reliably",
        "Localize (frontend vs backend vs data vs env)",
        "Inspect evidence (console, network tab, server logs, git diff)",
        "Form a hypothesis",
        "Change one thing",
        "Retest",
        "Keep or discard hypothesis",
    ])
    add_heading_styled(doc, "26.1 Browser DevTools minimum viable literacy", 2)
    bullets(doc, [
        "Console: JS errors and log output",
        "Network: failed API calls, status codes, response bodies",
        "Elements: inspect DOM and CSS that actually applied",
        "Application/Storage: localStorage, cookies (where relevant)",
    ])
    add_heading_styled(doc, "26.2 How to ask Claude for debugging help well", 2)
    p(doc, "Bad: “it’s broken.” Good: environment + steps + expected + actual + error text + what you already tried + “fix root cause, don’t hide the error.”")

    add_heading_styled(doc, "Part 27 — Security & Responsibility Module (Required Teaching)", 1)
    bullets(doc, [
        "Secrets management: .env, host env vars, rotation after leaks",
        "Never trust client-only checks for authorization",
        "Dependency caution: more packages = more supply chain risk",
        "User data minimization: collect only what you need",
        "Honest scope: school projects should not handle real payments/health data",
        "AI-generated code can include insecure defaults — run /security-review before deploy",
    ])

    add_heading_styled(doc, "Part 28 — Sample 90-minute Lesson Scripts (Teach Tomorrow)", 1)
    add_heading_styled(doc, "28.1 Lesson: Slash commands that matter", 2)
    table_data(doc, ["Minutes", "Activity"], [
        ["0–10", "Warm-up: students show hello page; error hospital"],
        ["10–25", "Teach context window + why /clear and /compact exist"],
        ["25–40", "Live demo /plan → approve → implement → /diff"],
        ["40–75", "Lab: tip calculator with forced plan mode"],
        ["75–85", "Share-outs: one edge case found"],
        ["85–90", "Homework: prompt rewrites + /help scavenger list"],
    ])
    add_heading_styled(doc, "28.2 Lesson: Git + GitHub first push", 2)
    table_data(doc, ["Minutes", "Activity"], [
        ["0–15", "Metaphors + live git status/add/commit"],
        ["15–30", "Create .gitignore; explain secrets"],
        ["30–50", "GitHub repo create + auth method (gh)"],
        ["50–75", "Students push; TA swarm on auth issues"],
        ["75–90", "README minimum sections written with Claude"],
    ])
    add_heading_styled(doc, "28.3 Lesson: HTML/CSS/JS as a system", 2)
    table_data(doc, ["Minutes", "Activity"], [
        ["0–15", "Three-layer model with live edit of one page"],
        ["15–35", "HTML structure workshop (semantic tags)"],
        ["35–55", "CSS box model + flex header"],
        ["55–80", "JS click counter state→render pattern"],
        ["80–90", "Map quiz: which layer is this bug in?"],
    ])

    page_break(doc)
    add_heading_styled(doc, "Appendix F — Full Slash Command Teaching Table (Core + Power)", 1)
    table_data(doc, ["Command", "Category", "Teach in week", "One-line student meaning"], [
        ["/help", "Meta", "1", "Show me what I can type"],
        ["/login", "Auth", "1", "Connect my Claude account"],
        ["/doctor", "Setup", "1", "Check if install is healthy"],
        ["/init", "Memory", "1–2", "Start project brain (CLAUDE.md)"],
        ["/plan", "Workflow", "2", "Think before coding"],
        ["/clear", "Context", "2", "New conversation"],
        ["/compact", "Context", "2–4", "Summarize to free space"],
        ["/diff", "Review", "2+", "Show what changed"],
        ["/rewind", "Recovery", "2+", "Undo bad direction"],
        ["/resume", "Session", "2+", "Continue old chat"],
        ["/permissions", "Safety", "3+", "What Claude may do"],
        ["/config", "Settings", "3+", "Preferences / model"],
        ["/model", "Settings", "4+", "Pick model"],
        ["/memory", "Memory", "4+", "Edit lasting instructions"],
        ["/context", "Context", "4+", "See context usage"],
        ["/usage", "Account", "any", "Usage awareness"],
        ["/code-review", "Quality", "7+", "Review my diff"],
        ["/security-review", "Security", "8+", "Find security issues"],
        ["/mcp", "Extensibility", "11+", "External tools"],
        ["/batch", "Scale", "advanced", "Huge parallel changes"],
        ["/debug", "Diagnostics", "as needed", "Debug Claude/session"],
        ["/btw", "UX", "any", "Side question"],
        ["/copy", "UX", "any", "Copy answer"],
        ["/desktop", "Surface", "optional", "Jump to desktop app"],
    ])

    add_heading_styled(doc, "Appendix G — Language Recognition Cheat Sheets", 1)
    p(doc, "HTML sniff test: lots of <tags>, structure of a document.")
    p(doc, "CSS sniff test: selectors and curly braces with property: value; pairs.")
    p(doc, "JavaScript sniff test: functions, const/let, addEventListener, fetch, logic.")
    p(doc, "Python sniff test: def, indentation blocks, colons at end of headers, pip/venv.")
    p(doc, "JSON sniff test: strict braces/brackets, double quotes on keys, used for data interchange.")

    add_heading_styled(doc, "Appendix H — Capstone PRD Template (Full Text)", 1)
    code_block(doc, [
        "App name:",
        "Problem statement:",
        "Primary user:",
        "Job to be done (1 sentence):",
        "MVP features (3-5):",
        "1.",
        "2.",
        "3.",
        "Out of scope:",
        "Screens:",
        "Data model (entities + fields):",
        "Auth needs (yes/no + why):",
        "Success demo script (60-120 seconds):",
        "Risks & unknowns:",
        "Milestone plan M1-M6:",
    ])

    doc.add_paragraph()
    callout(doc, "Final word to the instructor", [
        "You are teaching people to command a powerful agentic engineer and to understand the world that engineer operates in.",
        "Philosophy without mechanisms produces confidence without competence.",
        "Mechanisms without empathy produce fear.",
        "This guide aims at both: deep explanation and human pacing.",
        "Go slow in weeks 1–2. Demand verification always. Ship something real.",
        "Prefer this Edition 2 master guide over the thinner first draft.",
    ])

    p(doc, "— End of Master Teaching Guide (Edition 2.0 — Explanation-First) —", bold=True)
    p(doc, "Primary references: Claude Code docs at code.claude.com (commands, setup, best practices, skills, permission modes, memory). Always re-check /help for version-specific commands before live teaching.", size=9, italic=True, color=MUTED)

    doc.save(OUT)
    print("Wrote", OUT)
    print("Size KB:", round(os.path.getsize(OUT) / 1024, 1))

if __name__ == "__main__":
    build()
