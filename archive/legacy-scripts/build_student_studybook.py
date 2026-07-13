#!/usr/bin/env python3
"""
Student Study Book — explanation-first companion (not blank worksheets only)
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

OUT = "/home/shailesh/Downloads/Training/vibe-coding/Vibe_Coding_Student_Study_Book_v2.docx"
PRIMARY = RGBColor(0x0B, 0x3D, 0x2E)
ACCENT = RGBColor(0x1A, 0x7A, 0x5C)
DARK = RGBColor(0x1C, 0x28, 0x33)
MUTED = RGBColor(0x5D, 0x6D, 0x7E)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

def set_run(run, size=11, bold=False, italic=False, color=DARK, font="Calibri"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color

def H(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = PRIMARY if level == 1 else ACCENT if level == 2 else DARK
    return h

def p(doc, text, size=11, bold=False, italic=False, color=DARK, after=8):
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(after)
    para.paragraph_format.line_spacing = 1.15
    r = para.add_run(text)
    set_run(r, size=size, bold=bold, italic=italic, color=color)
    return para

def bullets(doc, items):
    for item in items:
        para = doc.add_paragraph(style="List Bullet")
        para.paragraph_format.space_after = Pt(3)
        para.clear()
        set_run(para.add_run(item), size=11)

def numbered(doc, items):
    for item in items:
        para = doc.add_paragraph(style="List Number")
        para.paragraph_format.space_after = Pt(3)
        para.clear()
        set_run(para.add_run(item), size=11)

def code(doc, lines):
    if isinstance(lines, str):
        lines = lines.split("\n")
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    cell._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="1E2A32"/>'))
    cell.text = ""
    first = True
    for line in lines:
        para = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        para.paragraph_format.space_after = Pt(0)
        set_run(para.add_run(line or " "), size=9, color=RGBColor(0xE8, 0xF0, 0xF2), font="Consolas")
    doc.add_paragraph()

def table_data(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        set_run(cell.paragraphs[0].add_run(h), size=10, bold=True, color=WHITE)
        cell._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="0B3D2E"/>'))
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            cell.text = ""
            set_run(cell.paragraphs[0].add_run(str(val)), size=9)
            if ri % 2 == 0:
                cell._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="F4F6F7"/>'))
    doc.add_paragraph()

def callout(doc, title, lines, fill="E8F6F1"):
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    cell._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill}"/>'))
    cell.text = ""
    set_run(cell.paragraphs[0].add_run(title), size=11, bold=True, color=PRIMARY)
    for line in lines:
        pp = cell.add_paragraph()
        set_run(pp.add_run(line), size=10)
    doc.add_paragraph()

def pb(doc):
    doc.add_page_break()

def practice(doc, title, questions):
    H(doc, title, 3)
    for i, q in enumerate(questions, 1):
        p(doc, f"{i}. {q}", bold=True)
        p(doc, "Answer: ________________________________________________", color=MUTED, size=10)
        p(doc, "_" * 70, color=MUTED, size=10)

def build():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.75)
    sec.bottom_margin = Inches(0.75)
    sec.left_margin = Inches(0.85)
    sec.right_margin = Inches(0.85)

    for _ in range(3):
        doc.add_paragraph()
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(t.add_run("STUDENT STUDY BOOK"), size=14, bold=True, color=ACCENT)
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(t.add_run("Vibe Coding with Claude Code"), size=26, bold=True, color=PRIMARY)
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(t.add_run("Explanations · Command Guides · Language Foundations · Labs"), size=12, italic=True, color=MUTED)
    doc.add_paragraph()
    callout(doc, "How to use this book", [
        "Read the explanation sections carefully — they are the real learning.",
        "Do every Try It lab on your laptop.",
        "Fill practice answers in your own words (not copied jargon).",
        "Use the Master Teaching Guide when you want even more depth or instructor context.",
        "If something fails: copy the error, use the bug template, then ask for help.",
    ])
    p(doc, "Name: _________________   Cohort: _________________   GitHub: _________________", size=10)
    pb(doc)

    # CH 1
    H(doc, "Chapter 1 — What You Are Learning")
    H(doc, "1.1 Vibe coding, explained carefully", 2)
    p(doc, "Vibe coding means you build software by describing goals in natural language while an AI coding agent (Claude Code) edits project files and runs tools. You are not “avoiding thinking.” You are moving your effort from typing syntax to deciding what to build, checking whether it works, and steering when the AI is wrong.")
    p(doc, "Claude Code is different from normal ChatGPT-style chat because it can act inside your project folder: read many files, change them, run install commands, run your app, and iterate based on the results.")
    H(doc, "1.2 Full-stack means the whole product stack", 2)
    bullets(doc, [
        "Frontend: what users see (HTML/CSS/JS in the browser)",
        "Backend: server logic (often Python or Node.js)",
        "Database/storage: lasting data",
        "Auth (often): login and “my data only”",
        "Deploy: public internet access",
    ])
    practice(doc, "Check your understanding", [
        "In your own words, what is the difference between Claude Code and a normal chatbot?",
        "Name the layers of a full-stack app and one example technology for each.",
        "Why is verification still required if AI writes the code?",
    ])
    pb(doc)

    # CH 2 terminal
    H(doc, "Chapter 2 — Terminal Guide (Study Version)")
    H(doc, "2.1 Why the terminal exists", 2)
    p(doc, "The terminal is a text remote control for your computer. Buttons are friendly; commands are precise. Developers use the terminal because it is scriptable, fast, and universal across tools (Git, Node, Python, Claude Code).")
    H(doc, "2.2 Commands you must understand", 2)
    table_data(doc, ["Command", "What it does", "When you use it"], [
        ["pwd", "Shows current folder path", "Before running tools — confirm location"],
        ["ls / dir / Get-ChildItem", "Lists files", "See what is in the folder"],
        ["cd name", "Enter a folder", "Move into your project"],
        ["cd ..", "Go up one folder", "Back out of a folder"],
        ["mkdir name", "Create folder", "Start a new project folder"],
        ["clear / cls", "Clear screen", "Reduce clutter (files stay)"],
        ["claude", "Start Claude Code", "Begin an AI coding session"],
        ["claude --version", "Show installed version", "Verify install"],
        ["claude doctor", "Diagnose setup issues", "When something is broken"],
        ["git status", "Show Git changes", "Before commit"],
    ])
    H(doc, "2.3 Anatomy of a command", 2)
    p(doc, "command + arguments + flags. Example: mkdir week1-hello creates a folder named week1-hello. Quotes are needed if names have spaces: cd \"My Projects\".")
    H(doc, "2.4 PATH and “command not found”", 2)
    p(doc, "PATH is a list of folders where the shell looks for programs. If Claude is installed but not on PATH, the shell says command not found. Fix: add the install directory (often ~/.local/bin) to PATH, then open a new terminal.")
    callout(doc, "Try It — Terminal lab", [
        "1) Open terminal. 2) pwd. 3) mkdir -p vibe-coding-course/week1-hello  4) cd into it  5) ls  6) Write what path you see.",
    ], fill="FEF9E7")
    practice(doc, "Practice", [
        "What is the difference between a file and a folder?",
        "What does cd .. do?",
        "Why must you start Claude Code from the correct project folder?",
    ])
    pb(doc)

    # CH 3 Claude
    H(doc, "Chapter 3 — Claude Code Features (Student Encyclopedia)")
    H(doc, "3.1 What Claude can do in a session", 2)
    bullets(doc, [
        "Read and search your project files",
        "Create and edit code/docs",
        "Run shell commands (install, test, start server)",
        "Follow CLAUDE.md project rules",
        "Use skills and (later) MCP tools",
        "Help with Git commits and pull requests",
        "Review diffs and security issues with skills/commands",
    ])
    H(doc, "3.2 Permission modes (you must know these)", 2)
    p(doc, "Permission modes control how often Claude asks before acting. In the CLI, Shift+Tab cycles modes.")
    table_data(doc, ["Mode", "Meaning for you", "Use when"], [
        ["Manual (default)", "Asks before many edits/commands", "Learning — safest default"],
        ["Plan", "Researches and plans; no source edits", "Before big features"],
        ["acceptEdits", "Auto-approves many file edits", "Only when you will review diffs after"],
        ["bypassPermissions", "Skips safety prompts", "Never on class laptops"],
    ])
    H(doc, "3.3 Context window", 2)
    p(doc, "Claude can only “see” a limited amount of conversation + files + logs at once. When the context fills, quality drops. Habits: one task per session when possible; /clear for new tasks; /compact for long same-task sessions; keep CLAUDE.md short.")
    H(doc, "3.4 CLAUDE.md", 2)
    p(doc, "A markdown file Claude reads every session. Put run commands, project rules, and gotchas. Do not put essays. Create starter with /init, then edit.")
    code(doc, [
        "# Commands",
        "- Run: npm run dev",
        "# Rules",
        "- Never commit .env",
        "- Explain changes simply after big edits",
    ])
    H(doc, "3.5 Skills", 2)
    p(doc, "Skills are reusable instruction packs in SKILL.md files. They load when needed (or via /skill-name). Use them for repeated procedures (deploy checklist, beginner explanations). CLAUDE.md = always-on short rules. Skills = on-demand procedures.")
    H(doc, "3.6 MCP (intro only)", 2)
    p(doc, "MCP (Model Context Protocol) connects Claude to external tools (browser, databases, APIs). Powerful and advanced. Master local building first.")
    pb(doc)

    # CH 4 slash commands
    H(doc, "Chapter 4 — Slash Commands You Will Actually Use")
    p(doc, "Slash commands start with /. They control Claude Code. Below is your working encyclopedia.")

    commands = [
        ("/help", "Show help", "Anytime you forget features", "/help"),
        ("/login", "Authenticate account", "First run or auth errors", "/login"),
        ("/doctor", "Diagnose install/config", "Broken environment", "/doctor"),
        ("/init", "Create starter CLAUDE.md", "New project", "/init"),
        ("/plan", "Enter plan mode / plan a task", "Before coding features", "/plan Add dark mode. No code yet."),
        ("/clear", "New empty conversation", "Switching tasks", "/clear"),
        ("/compact", "Summarize to free context", "Long session, same task", "/compact Keep schema decisions"),
        ("/diff", "Review changes", "Before commit", "/diff"),
        ("/rewind", "Undo bad direction", "Claude made a mess", "/rewind"),
        ("/resume", "Open past sessions", "Continue later", "/resume"),
        ("/permissions", "Tool allow/ask/deny rules", "Prompt fatigue or lockdown", "/permissions"),
        ("/config", "Settings UI / key=value", "Model, theme, prefs", "/config"),
        ("/memory", "Refine memory/instructions", "Claude ignores project rules", "/memory"),
        ("/context", "See context usage", "Quality dropped", "/context"),
        ("/usage or /cost", "Usage info", "Monitor limits", "/usage"),
        ("/code-review", "Review diff for bugs/cleanups", "Before demo/merge", "/code-review"),
        ("/security-review", "Security-focused review", "Before deploy with users", "/security-review"),
        ("/mcp", "Manage MCP servers", "After MCP setup", "/mcp"),
        ("/copy", "Copy last response", "Save notes/snippets", "/copy"),
        ("/btw", "Side question", "Quick fact mid-task", "/btw what is CORS?"),
    ]
    for name, what, when, ex in commands:
        H(doc, name, 3)
        p(doc, f"What it does: {what}")
        p(doc, f"When: {when}")
        p(doc, f"Example: {ex}", italic=True)
    callout(doc, "Try It — Command drill", [
        "Start Claude in an empty folder. Run /help. Run /init. Ask Claude to create a hello HTML page.",
        "Use /plan for a tip calculator, approve, then implement. Use /diff before finishing.",
    ], fill="FEF9E7")
    practice(doc, "Practice", [
        "When should you use /plan instead of immediately asking for code?",
        "What is the difference between /clear and /compact?",
        "Why is /diff important before git commit?",
    ])
    pb(doc)

    # CH 5 prompt craft
    H(doc, "Chapter 5 — Prompt Craft That Works")
    H(doc, "5.1 PROMPT formula", 2)
    table_data(doc, ["Letter", "Meaning", "Your notes"], [
        ["P", "Purpose — outcome", ""],
        ["R", "Rules — constraints", ""],
        ["O", "Objects — screens/data/buttons", ""],
        ["M", "Musts — acceptance checks", ""],
        ["P", "Proof — how to verify", ""],
        ["T", "Tone/style", ""],
    ])
    H(doc, "5.2 Bug report template (memorize)", 2)
    code(doc, [
        "Bug report:",
        "Steps:",
        "Expected:",
        "Actual:",
        "Error text:",
        "Please find root cause, fix, verify, explain simply.",
    ])
    H(doc, "5.3 Rewrite drill", 2)
    p(doc, "Rewrite each bad prompt into a good one in your notebook:")
    bullets(doc, [
        "make a nice app",
        "fix everything",
        "add database",
        "make UI modern",
    ])
    pb(doc)

    # HTML
    H(doc, "Chapter 6 — HTML Study Guide")
    H(doc, "6.1 Definition", 2)
    p(doc, "HTML (HyperText Markup Language) labels the structure of a webpage: headings, paragraphs, links, images, forms, buttons. It is markup, not a full programming language. Browsers read HTML to build the page skeleton.")
    H(doc, "6.2 Element anatomy", 2)
    code(doc, '<a href="https://example.com" class="link">Click me</a>')
    bullets(doc, [
        "Tag name: a (anchor/link)",
        "Attributes: href (destination), class (CSS hook)",
        "Content: Click me",
        "Closing tag: </a>",
    ])
    H(doc, "6.3 Page skeleton you must recognize", 2)
    code(doc, [
        "<!DOCTYPE html>",
        '<html lang="en">',
        "<head> ... title, css links, meta ... </head>",
        "<body> ... visible content ... scripts ... </body>",
        "</html>",
    ])
    H(doc, "6.4 Tag map", 2)
    table_data(doc, ["Tag", "Role", "Remember"], [
        ["h1-h6", "Headings", "h1 = main title"],
        ["p", "Paragraph", "Body text"],
        ["a", "Link", "needs href"],
        ["img", "Image", "src + alt"],
        ["button", "Button", "pair with JS or form"],
        ["form/input/label", "Data entry", "labels matter"],
        ["ul/ol/li", "Lists", "nav and content lists"],
        ["header/main/footer/nav", "Semantic sections", "better than only div"],
        ["div/span", "Generic boxes", "use when no better tag"],
    ])
    H(doc, "6.5 Forms", 2)
    p(doc, "Forms collect user input. type controls keyboard and validation (email, password, number). required blocks empty submit in the browser. Labels improve accessibility and usability.")
    callout(doc, "Try It — HTML", [
        "Ask Claude Code to create index.html explaining each tag in comments.",
        "Open it in a browser. Change the title and an h1 yourself (or ask Claude) and reload.",
    ], fill="FEF9E7")
    practice(doc, "HTML practice", [
        "What goes in <head> vs <body>?",
        "Why does <img> need alt text?",
        "What is the difference between a div and a main tag?",
    ])
    pb(doc)

    # CSS
    H(doc, "Chapter 7 — CSS Study Guide")
    H(doc, "7.1 Definition", 2)
    p(doc, "CSS (Cascading Style Sheets) controls presentation: colors, fonts, spacing, layout, responsive design. HTML says what something is; CSS says how it looks.")
    H(doc, "7.2 Rule structure", 2)
    code(doc, [
        "selector {",
        "  property: value;",
        "}",
        "",
        "p { color: #222; font-size: 16px; }",
        ".card { padding: 16px; border-radius: 8px; }",
        "#logo { width: 48px; }",
    ])
    H(doc, "7.3 Box model", 2)
    p(doc, "Every element is a box: content + padding + border + margin. Layout bugs are often spacing bugs. When Claude “fixes spacing,” it is usually adjusting these.")
    H(doc, "7.4 Layout words you will hear", 2)
    bullets(doc, [
        "Flexbox: arrange items in a row or column easily",
        "Grid: two-dimensional page layouts",
        "Media queries: different styles for mobile vs desktop",
        "Responsive: usable across screen sizes",
    ])
    H(doc, "7.5 Quality bar", 2)
    bullets(doc, [
        "Readable font size (~16px body)",
        "Good contrast",
        "Consistent spacing",
        "Obvious primary button",
        "No accidental horizontal page scroll on mobile",
    ])
    practice(doc, "CSS practice", [
        "What is the difference between margin and padding?",
        "What does a class selector look like?",
        "Name two ways to make a page mobile-friendly.",
    ])
    pb(doc)

    # JS
    H(doc, "Chapter 8 — JavaScript Study Guide")
    H(doc, "8.1 Definition", 2)
    p(doc, "JavaScript is a programming language that makes pages interactive. It runs in the browser (and on servers via Node.js). Clicks, form checks, updating text without reload, calling APIs — this is JavaScript territory.")
    H(doc, "8.2 Core idea map", 2)
    table_data(doc, ["Idea", "Meaning", "Why apps need it"], [
        ["Variable", "Named value storage", "Remember score, email, etc."],
        ["Function", "Reusable actions", "Handle click / save"],
        ["if/else", "Decisions", "Show errors when invalid"],
        ["Array/object", "Lists and records", "Tasks, users, products"],
        ["DOM", "JS view of HTML", "Change page content"],
        ["Event", "User/system action", "click, submit, input"],
        ["fetch/API call", "Talk to backend", "Load/save real data"],
    ])
    H(doc, "8.3 Events (the interactive core)", 2)
    p(doc, "You attach a listener: when the user clicks Save, run a function. That function may validate input, update the page, and send data to a server.")
    code(doc, [
        "button.addEventListener(\"click\", () => {",
        "  // do something",
        "});",
    ])
    H(doc, "8.4 Node.js and npm (recognition level)", 2)
    p(doc, "Node.js runs JavaScript outside the browser. npm installs packages listed in package.json. npm install downloads dependencies. npm run dev often starts a development server. Claude may run these; you should understand the words.")
    practice(doc, "JavaScript practice", [
        "HTML/CSS/JS — which is structure, style, behavior?",
        "What is an event listener in plain language?",
        "What does npm install do?",
    ])
    pb(doc)

    # Python
    H(doc, "Chapter 9 — Python Study Guide")
    H(doc, "9.1 Why Python", 2)
    p(doc, "Python is a readable general-purpose language widely used for backends, automation, and AI tooling. Frameworks like Flask and FastAPI turn Python functions into web APIs.")
    H(doc, "9.2 Reading Python", 2)
    code(doc, [
        "def add(a, b):",
        "    return a + b",
        "",
        "print(add(2, 3))  # 5",
    ])
    bullets(doc, [
        "def defines a function",
        "Indentation defines blocks (spaces matter)",
        "print writes to the terminal",
        "Run: python3 file.py",
    ])
    H(doc, "9.3 Backend mental model", 2)
    numbered(doc, [
        "Browser requests /api/tasks",
        "Python route/function runs",
        "Database read/write happens",
        "JSON response returns",
        "Frontend renders data",
    ])
    H(doc, "9.4 pip and virtual environments", 2)
    p(doc, "pip installs Python packages. A virtual environment (venv) isolates dependencies per project so packages do not collide. requirements.txt lists packages so others can recreate the environment.")
    practice(doc, "Python practice", [
        "Why does indentation matter in Python?",
        "What is a web route in plain language?",
        "Why use a virtual environment?",
    ])
    pb(doc)

    # Git
    H(doc, "Chapter 10 — Git Study Guide")
    H(doc, "10.1 Problem Git solves", 2)
    p(doc, "Git stores snapshots of your project over time. You can undo disasters, understand history, and collaborate. It is not the same thing as GitHub (GitHub hosts Git repos online).")
    H(doc, "10.2 Essential terms", 2)
    table_data(doc, ["Term", "Meaning"], [
        ["Repository", "Project tracked by Git"],
        ["Commit", "Snapshot + message"],
        ["Stage (add)", "Mark changes for next commit"],
        ["Branch", "Parallel line of work"],
        ["Remote", "Hosted copy (often GitHub)"],
        ["Push/Pull", "Upload / download commits"],
        ["Clone", "Copy repo to your machine"],
        [".gitignore", "Patterns Git should not track"],
    ])
    H(doc, "10.3 Daily commands", 2)
    code(doc, [
        "git status",
        "git add .",
        "git commit -m \"Describe why this change exists\"",
        "git push",
        "git pull",
        "git log --oneline",
    ])
    H(doc, "10.4 Secrets rule", 2)
    callout(doc, "Never commit .env or API keys", [
        "If you commit a secret, revoke it immediately and create a new key.",
        "Check git status before every push.",
    ], fill="FDEDEC")
    practice(doc, "Git practice", [
        "What is the difference between git and GitHub?",
        "What makes a good commit message?",
        "Why is .gitignore important?",
    ])
    pb(doc)

    # GitHub
    H(doc, "Chapter 11 — GitHub Study Guide")
    H(doc, "11.1 What GitHub adds", 2)
    bullets(doc, [
        "Cloud backup of your Git repo",
        "Portfolio of your work",
        "Collaboration: issues, pull requests, reviews",
        "Automation via GitHub Actions (later topic)",
    ])
    H(doc, "11.2 Setup checklist", 2)
    numbered(doc, [
        "Create account + verify email",
        "Enable 2FA",
        "Create a repository with README + .gitignore",
        "Authenticate locally (gh auth login OR token/SSH)",
        "Push your project",
        "Write a real README (features, setup, run, demo URL)",
    ])
    H(doc, "11.3 Pull Request concept", 2)
    p(doc, "A Pull Request (PR) proposes merging changes from one branch into another. Even solo, PRs teach professional review habits. Claude can help write PR summaries and run /code-review.")
    H(doc, "11.4 README quality checklist", 2)
    bullets(doc, [
        "What the app does",
        "Screenshot",
        "Features",
        "Stack",
        "Install/run steps",
        "Environment variables needed (names, not secret values)",
        "Live URL",
        "Your name",
    ])
    callout(doc, "Try It — GitHub", [
        "Put Project 1 or 2 on GitHub with a solid README.",
        "Ask Claude: ensure .env ignored; write README; commit; guide push.",
    ], fill="FEF9E7")
    practice(doc, "GitHub practice", [
        "Why enable 2FA?",
        "What should never appear in a public repo?",
        "List 5 sections of a strong README.",
    ])
    pb(doc)

    # Full stack path
    H(doc, "Chapter 12 — Full-Stack Build Path (Your Labs)")
    table_data(doc, ["Project", "You learn", "Done?"], [
        ["0 Hello page", "Install, folders, Claude basics", ""],
        ["1 Portfolio site", "HTML/CSS multi-page", ""],
        ["2 Interactive tool", "JS events/validation", ""],
        ["3 Persistent notes/todo", "CRUD + storage", ""],
        ["4 API + UI", "Backend bridge", ""],
        ["5 Auth", "Users & ownership", ""],
        ["6 Capstone", "PRD → ship → demo", ""],
    ])
    H(doc, "12.1 Milestone rule", 2)
    numbered(doc, [
        "/plan the milestone",
        "Approve plan",
        "Implement only that milestone",
        "Verify yourself in browser/terminal",
        "Commit",
        "Next milestone",
    ])
    H(doc, "12.2 Capstone minimum", 2)
    bullets(doc, [
        "Real problem + user",
        "PRD with 3–5 MVP features",
        "Working CRUD + persistence",
        "GitHub + README",
        "Deployed URL",
        "5-minute demo with 3 human decisions explained",
    ])
    pb(doc)

    # Deploy
    H(doc, "Chapter 13 — Deploy Study Guide")
    p(doc, "Deploy means hosting your app on the public internet. Localhost is only your computer. Production needs build/start commands, environment variables, and often a hosted database.")
    p(doc, "Common failure: app works locally because secrets exist in your local .env, but production has no env vars set. Always configure them in the host dashboard.")
    numbered(doc, [
        "Confirm local run is clean",
        "Set env vars on host",
        "Deploy",
        "Smoke test in incognito + phone",
        "Add URL to README",
    ])
    pb(doc)

    # Reference
    H(doc, "Chapter 14 — Quick Reference Cards")
    H(doc, "14.1 Stuck protocol", 2)
    numbered(doc, [
        "State what you tried and expected",
        "Copy exact error",
        "Paste bug template to Claude",
        "If still stuck after one careful loop → ask human help",
    ])
    H(doc, "14.2 Daily builder checklist", 2)
    bullets(doc, [
        "Correct folder",
        "Claude started there",
        "Purpose + musts + proof in prompt",
        "Plan mode for big work",
        "I ran the app myself",
        "Tried an edge case",
        "Committed working progress",
        "No secrets staged",
    ])
    H(doc, "14.3 Essential slash commands", 2)
    code(doc, "/help /login /doctor /init /plan /clear /compact /diff /rewind /resume /code-review /security-review")
    H(doc, "14.4 Essential Git", 2)
    code(doc, "status → add → commit → push  |  pull to update  |  .gitignore protects secrets")

    doc.add_paragraph()
    callout(doc, "You are becoming a builder", [
        "If you can explain the parts, direct Claude Code with precision, verify results,",
        "use Git/GitHub responsibly, and ship a URL — you have real power.",
        "Keep this book. Re-read chapters when a project gets hard.",
    ])
    p(doc, "— End of Student Study Book —", bold=True)

    doc.save(OUT)
    print("Wrote", OUT, "KB", round(os.path.getsize(OUT)/1024, 1))

if __name__ == "__main__":
    build()
