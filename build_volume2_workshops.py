#!/usr/bin/env python3
"""Volume 2 — Deep Workshops: HTML/CSS/JS/Python/GitHub/Claude step-by-step with explanations."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

OUT = "/home/shailesh/Downloads/Training/vibe-coding/Vibe_Coding_Deep_Workshops_v2.docx"
PRIMARY = RGBColor(0x0B, 0x3D, 0x2E)
ACCENT = RGBColor(0x1A, 0x7A, 0x5C)
DARK = RGBColor(0x1C, 0x28, 0x33)
MUTED = RGBColor(0x5D, 0x6D, 0x7E)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

def sr(run, size=11, bold=False, italic=False, color=DARK, font="Calibri"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size); run.bold = bold; run.italic = italic; run.font.color.rgb = color

def H(doc, t, level=1):
    h = doc.add_heading(t, level=level)
    for r in h.runs:
        r.font.color.rgb = PRIMARY if level == 1 else ACCENT if level == 2 else DARK

def p(doc, t, size=11, bold=False, italic=False, color=DARK):
    para = doc.add_paragraph(); para.paragraph_format.space_after = Pt(8); para.paragraph_format.line_spacing = 1.15
    sr(para.add_run(t), size=size, bold=bold, italic=italic, color=color)

def bullets(doc, items):
    for i in items:
        para = doc.add_paragraph(style="List Bullet"); para.paragraph_format.space_after = Pt(3); para.clear()
        sr(para.add_run(i))

def numbered(doc, items):
    for i in items:
        para = doc.add_paragraph(style="List Number"); para.paragraph_format.space_after = Pt(3); para.clear()
        sr(para.add_run(i))

def code(doc, lines):
    if isinstance(lines, str): lines = lines.split("\n")
    table = doc.add_table(rows=1, cols=1); cell = table.rows[0].cells[0]
    cell._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="1E2A32"/>')); cell.text = ""
    first = True
    for line in lines:
        para = cell.paragraphs[0] if first else cell.add_paragraph(); first = False
        para.paragraph_format.space_after = Pt(0)
        sr(para.add_run(line or " "), size=9, color=RGBColor(0xE8,0xF0,0xF2), font="Consolas")
    doc.add_paragraph()

def table_data(doc, headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers)); t.style = "Table Grid"
    for i,h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = ""; sr(c.paragraphs[0].add_run(h), size=10, bold=True, color=WHITE)
        c._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="0B3D2E"/>'))
    for ri,row in enumerate(rows):
        for ci,val in enumerate(row):
            c = t.rows[ri+1].cells[ci]; c.text = ""; sr(c.paragraphs[0].add_run(str(val)), size=9)
            if ri%2==0: c._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="F4F6F7"/>'))
    doc.add_paragraph()

def callout(doc, title, lines, fill="E8F6F1"):
    table = doc.add_table(rows=1, cols=1); cell = table.rows[0].cells[0]
    cell._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill}"/>')); cell.text = ""
    sr(cell.paragraphs[0].add_run(title), size=11, bold=True, color=PRIMARY)
    for line in lines:
        pp = cell.add_paragraph(); sr(pp.add_run(line), size=10)
    doc.add_paragraph()

def pb(doc): doc.add_page_break()

def build():
    doc = Document()
    s = doc.sections[0]
    s.top_margin = Inches(0.75); s.bottom_margin = Inches(0.75)
    s.left_margin = Inches(0.85); s.right_margin = Inches(0.85)

    for _ in range(3): doc.add_paragraph()
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr(t.add_run("VOLUME 2 — DEEP WORKSHOPS"), size=14, bold=True, color=ACCENT)
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr(t.add_run("Vibe Coding with Claude Code"), size=26, bold=True, color=PRIMARY)
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr(t.add_run("Step-by-step labs with full explanations (HTML, CSS, JS, Python, GitHub, Claude Code)"), size=11, italic=True, color=MUTED)
    doc.add_paragraph()
    callout(doc, "How to use Volume 2", [
        "This volume is practice with teaching commentary — not empty slogans.",
        "Do workshops in order. Read every “Why this works” box.",
        "When Claude writes code, pause and map each file to the concept it demonstrates.",
        "Pair with Master Teaching Guide (theory) and Student Study Book (review questions).",
    ])
    pb(doc)

    # ========== WS1 Claude install deep ==========
    H(doc, "Workshop 1 — Install Claude Code and Understand Every Step")
    H(doc, "Goal", 2)
    p(doc, "Leave with a working Claude Code install and the ability to explain what each install step did.")
    H(doc, "Step-by-step (macOS / Linux)", 2)
    numbered(doc, [
        "Open Terminal.",
        "Run the official install script: curl -fsSL https://claude.ai/install.sh | bash",
        "What this does: downloads Anthropic’s installer and runs it. The pipe | sends download output into bash.",
        "If command not found afterward: your shell cannot find the claude program. Add ~/.local/bin to PATH (installer notes usually show exact lines for zsh/bash), then open a NEW terminal.",
        "Run claude --version — should print a version number.",
        "Run claude doctor — health check for install/config issues.",
        "mkdir -p ~/vibe-coding-course/ws1 && cd ~/vibe-coding-course/ws1",
        "Run claude — browser login appears. Use a plan that includes Claude Code (Pro/Max/Team/etc.).",
        "Inside Claude, type /help and skim available commands.",
    ])
    H(doc, "Windows notes explained", 2)
    bullets(doc, [
        "PowerShell uses: irm https://claude.ai/install.ps1 | iex  (irm = Invoke-RestMethod, iex = Invoke-Expression).",
        "If you see irm not recognized, you are probably in CMD, not PowerShell.",
        "If you see && is not valid, you pasted a CMD-style chain into PowerShell.",
        "Install Git for Windows so Claude can use Git Bash tooling more smoothly.",
    ])
    H(doc, "First prompt — and what Claude does behind the scenes", 2)
    code(doc, "Create a file hello.html that says Hello from Workshop 1 with a friendly style. Then tell me how to open it.")
    p(doc, "Behind the scenes, Claude will typically: (1) decide to write a file, (2) request permission if in Manual mode, (3) write HTML/CSS, (4) explain how to open it. Your job: approve carefully, open the file in a browser, confirm it works.")
    callout(doc, "Checkpoint", [
        "I can open terminal and navigate to a project folder.",
        "claude --version works.",
        "I generated and opened hello.html.",
        "I can explain PATH and why command-not-found happens.",
    ], fill="FEF9E7")
    pb(doc)

    # ========== WS2 Slash command lab ==========
    H(doc, "Workshop 2 — Slash Commands Lab (Do / Observe / Explain)")
    p(doc, "For each command: run it, write what you saw, explain when you would use it again.")
    table_data(doc, ["Command", "What I observed", "I will use it when…"], [
        ["/help", "", ""],
        ["/init", "", ""],
        ["/plan (with a tiny task)", "", ""],
        ["/clear", "", ""],
        ["/compact (after a long chat)", "", ""],
        ["/diff (after edits)", "", ""],
        ["/context", "", ""],
        ["/doctor", "", ""],
    ])
    H(doc, "Guided scenario", 2)
    numbered(doc, [
        "In a new folder, /init — open the created CLAUDE.md and rewrite it shorter in your own words (with Claude’s help).",
        "/plan Build a tip calculator single HTML file with validation. No code yet.",
        "Approve or correct the plan in English.",
        "Implement only the approved plan.",
        "/diff — describe three changes Claude made.",
        "Manually test: empty input, negative bill, normal 15% tip.",
        "Ask Claude to commit with a good message (if Git is ready) OR just save files.",
    ])
    H(doc, "Why plan mode is not optional for quality", 2)
    p(doc, "Without planning, Claude optimizes for “produce code now.” With planning, you catch wrong assumptions while they are free. Professional builders plan; beginners need planning even more because they cannot instantly spot architectural mistakes.")
    pb(doc)

    # ========== WS3 HTML deep ==========
    H(doc, "Workshop 3 — HTML Deep Build (Portfolio Page)")
    H(doc, "Concepts you must be able to explain after this workshop", 2)
    bullets(doc, [
        "Difference between head and body",
        "What an attribute is (href, src, alt, class, id)",
        "Semantic tags vs div soup",
        "How a form field becomes data later",
        "Why accessibility (labels, alt) is part of quality, not decoration",
    ])
    H(doc, "Build with Claude — but read the output", 2)
    code(doc, [
        "/plan Create a one-page personal portfolio in HTML only first",
        "(inline or linked CSS later). Sections: hero, about, skills list,",
        "projects (2 cards), contact form (visual). Use semantic tags.",
        "After building, produce a table: each tag used → purpose in plain English.",
        "Do not use a JS framework.",
    ])
    H(doc, "Teacher commentary on common tags (study this)", 2)
    table_data(doc, ["You will see", "What it means", "Common mistake"], [
        ["<header>", "Top branding/nav region", "Putting entire page inside header"],
        ["<nav>", "Navigation links", "Using div with no nav semantics"],
        ["<main>", "Primary unique content", "Multiple mains"],
        ["<section>", "Thematic grouping", "Meaningless wrapper with no heading"],
        ["<article>", "Self-contained piece", "Using for tiny decorative bits"],
        ["<h1>", "Page’s main heading", "Many h1s competing"],
        ['<a href="">', "Hyperlink", "href missing or # everywhere"],
        ['<img alt="">', "Image with text alternative", "Empty alt on meaningful images"],
        ["<label for>", "Names a control", "Placeholder used instead of label"],
        ['<input type="email">', "Email field", "type=text for emails always"],
    ])
    H(doc, "Exercise: break and fix", 2)
    numbered(doc, [
        "Remove an alt attribute; ask Claude what accessibility problem you created.",
        "Change an h1 to a div styled large — discuss why semantics still matter.",
        "Add a second page about.html and link between pages with relative paths (about.html not full https URL).",
    ])
    p(doc, "Relative path means “relative to this file’s location.” If index.html and about.html are in the same folder, href=\"about.html\" works. Nested folders need paths like projects/app.html or ../index.html (up one folder).")
    pb(doc)

    # ========== WS4 CSS deep ==========
    H(doc, "Workshop 4 — CSS Deep Build (Make It Professional)")
    H(doc, "Concepts", 2)
    bullets(doc, [
        "Selectors: element, class, id",
        "Box model: content, padding, border, margin",
        "Flexbox for alignment",
        "Responsive rules with max-width and media queries",
        "Variables (custom properties) for colors",
    ])
    H(doc, "Prompt", 2)
    code(doc, [
        "Create styles.css for my portfolio. Requirements:",
        "- CSS variables for colors and spacing",
        "- Body font 16px+, line-height 1.5",
        "- Header with flexbox space-between",
        "- Project cards in a responsive grid",
        "- Mobile: stack navigation",
        "Then explain box model using one card as example.",
        "Do not change HTML structure unless required.",
    ])
    H(doc, "Read this CSS like a teacher", 2)
    code(doc, [
        ":root {",
        "  --bg: #f7f3eb;",
        "  --text: #1a2421;",
        "  --brand: #1a7a5c;",
        "  --space: 16px;",
        "}",
        "body {",
        "  margin: 0;",
        "  font-family: system-ui, sans-serif;",
        "  color: var(--text);",
        "  background: var(--bg);",
        "}",
        ".card {",
        "  padding: var(--space);",
        "  margin-bottom: var(--space);",
        "  border-radius: 12px;",
        "  background: white;",
        "}",
    ])
    bullets(doc, [
        ":root defines reusable variables — change brand color once, update everywhere.",
        "margin: 0 on body removes default browser gaps.",
        "padding inside .card creates inner breathing room.",
        "margin-bottom creates outer separation between cards.",
    ])
    H(doc, "Flexbox micro-lesson", 2)
    p(doc, "display: flex turns a container into a flex formatting context. justify-content controls main axis distribution (e.g., space-between). align-items controls cross-axis alignment (e.g., center). This is how modern navbars put logo left and links right without fragile floats.")
    H(doc, "Responsive micro-lesson", 2)
    code(doc, [
        "@media (max-width: 600px) {",
        "  nav { flex-direction: column; }",
        "}",
    ])
    p(doc, "When the viewport is 600px wide or less, apply these rules. That is the heart of responsive CSS.")
    pb(doc)

    # ========== WS5 JS deep ==========
    H(doc, "Workshop 5 — JavaScript Deep Build (Quiz App)")
    H(doc, "Concepts", 2)
    bullets(doc, [
        "Variables and arrays of objects (question bank)",
        "Functions",
        "DOM selection and text updates",
        "Click events",
        "Simple state: current index, score",
        "Conditional rendering (show results when finished)",
    ])
    H(doc, "State — the most important idea", 2)
    p(doc, "State means “the data that describes the app right now.” For a quiz: which question number are we on? what is the score? did we finish? UI is a projection of state. When state changes, UI should update. Bugs often mean state and UI disagree.")
    H(doc, "Prompt", 2)
    code(doc, [
        "/plan Build a beginner quiz app with HTML/CSS/JS files.",
        "5 questions as an array of objects: {question, choices[4], answerIndex}.",
        "Show one question at a time, track score, restart button.",
        "Validate that users must pick an answer before next.",
        "After implementation, walk me through the state variables",
        "and the click handler step by step like I am new.",
    ])
    H(doc, "Map code regions (fill while reading Claude’s code)", 2)
    table_data(doc, ["Region", "File/location", "Job"], [
        ["Question data", "", "Stores Q&A content"],
        ["State variables", "", "Tracks progress"],
        ["render() or equivalent", "", "Updates DOM from state"],
        ["Next/Submit handler", "", "Changes state on click"],
        ["Restart handler", "", "Resets state"],
    ])
    H(doc, "Manual tests you must run", 2)
    numbered(doc, [
        "Answer all correctly → perfect score",
        "Answer all wrong → zero",
        "Try next without selecting → blocked or message",
        "Restart returns to question 1 with score 0",
        "Mobile width still usable",
    ])
    pb(doc)

    # ========== WS6 Python ==========
    H(doc, "Workshop 6 — Python Backend Mini API")
    H(doc, "Why this workshop exists", 2)
    p(doc, "Frontends alone cannot safely be the only source of truth for multi-user data. A backend receives requests, applies rules, talks to storage, returns JSON. You will build a tiny tasks API and call it from a simple page (or use tools like curl).")
    H(doc, "Concepts", 2)
    bullets(doc, [
        "Python functions and modules",
        "Virtual environment + pip",
        "HTTP methods: GET (read), POST (create), PUT/PATCH (update), DELETE (remove)",
        "Routes/endpoints like /api/tasks",
        "JSON request and response bodies",
        "Status codes: 200 OK, 201 created, 400 bad request, 404 not found",
    ])
    H(doc, "Prompt (FastAPI or Flask — pick one stack for class)", 2)
    code(doc, [
        "/plan Create a beginner task API in Python:",
        "- in-memory list is OK for first version (explain tradeoffs)",
        "- GET /api/tasks list",
        "- POST /api/tasks create {title}",
        "- DELETE /api/tasks/{id}",
        "- Input validation for empty title",
        "Include venv + requirements.txt + run instructions.",
        "Then add a minimal HTML page that fetches and displays tasks.",
        "Explain each endpoint like a teacher.",
    ])
    H(doc, "How a request flows", 2)
    numbered(doc, [
        "Browser/JS sends GET /api/tasks",
        "Server framework matches the route",
        "Your function runs",
        "Returns JSON list",
        "JS parses JSON and updates the DOM",
    ])
    H(doc, "In-memory vs database", 2)
    p(doc, "In-memory storage dies when the server restarts. Databases persist. Workshop path: memory first (learn API shape), then SQLite/Postgres persistence in the next workshop.")
    callout(doc, "Security seed", [
        "Never trust frontend validation alone.",
        "Backend must validate again.",
        "Later with auth: always check which user owns a task before delete/update.",
    ], fill="FDEDEC")
    pb(doc)

    # ========== WS7 persistence ==========
    H(doc, "Workshop 7 — Persistence & CRUD Properly")
    p(doc, "CRUD is the backbone of business apps. Implement create/read/update/delete with data that survives restart (SQLite is an excellent class choice: a single file database, low ops burden).")
    code(doc, [
        "/plan Upgrade the task API to SQLite persistence.",
        "Schema: id, title, done, created_at.",
        "Migrations or startup schema creation explained simply.",
        "Provide sample curl commands for each CRUD operation.",
        "Do not add auth yet.",
    ])
    H(doc, "SQL ideas in plain language", 2)
    table_data(doc, ["Idea", "Meaning"], [
        ["Table", "Spreadsheet-like structure for one entity (tasks)"],
        ["Row", "One task"],
        ["Column", "Field (title, done)"],
        ["Primary key", "Unique id for each row"],
        ["Query", "Question to the database (SELECT/INSERT/UPDATE/DELETE)"],
    ])
    p(doc, "You do not need to become a DBA. You need to recognize that storage is structured and that Claude’s SQL should match your schema.")
    pb(doc)

    # ========== WS8 Git ==========
    H(doc, "Workshop 8 — Git From Zero to Daily Habit")
    H(doc, "Initialize and first commit", 2)
    code(doc, [
        "git init",
        "git status",
        "# create .gitignore first (node_modules, venv, .env)",
        "git add .",
        "git status   # staged vs unstaged understanding",
        "git commit -m \"Initial project scaffold\"",
        "git log --oneline",
    ])
    H(doc, "What each command really does", 2)
    bullets(doc, [
        "git init: creates a hidden .git directory that stores history.",
        "git status: compares working files vs last commit vs staging area.",
        "git add: marks files to include in the next snapshot.",
        "git commit: writes a snapshot with metadata (message, time, author).",
        "git log: reads history backwards.",
    ])
    H(doc, "Using Claude safely with Git", 2)
    code(doc, [
        "Show me git status and explain staged vs unstaged in plain English.",
        "Draft a commit message for these changes.",
        "Commit only the intended files. Do not push yet.",
        "/diff",
    ])
    H(doc, "Recovery literacy", 2)
    p(doc, "If you commit too early: you can add another commit (preferred beginner path) rather than rewriting history. If Claude messes files: /rewind or git checkout/restore carefully (ask for explanation before destructive resets). Avoid git reset --hard until you understand it deletes uncommitted work.")
    pb(doc)

    # ========== WS9 GitHub ==========
    H(doc, "Workshop 9 — GitHub Complete Hands-On")
    H(doc, "A. Account hardening", 2)
    numbered(doc, [
        "Create GitHub account with professional username",
        "Verify email",
        "Enable 2FA",
        "Optional: profile README (later)",
    ])
    H(doc, "B. Create remote repo", 2)
    numbered(doc, [
        "New repository → name it after your project",
        "Add README optional if local already has one (avoid merge confusion; beginners often create empty repo without README if local exists)",
        "Copy remote URL (HTTPS or SSH)",
    ])
    H(doc, "C. Connect local → GitHub", 2)
    code(doc, [
        "git remote add origin https://github.com/USER/REPO.git",
        "git branch -M main",
        "git push -u origin main",
    ])
    p(doc, "Authentication options explained:")
    bullets(doc, [
        "GitHub CLI: gh auth login — guided, great for class",
        "HTTPS + personal access token — token acts like a password for Git",
        "SSH keys — generate key pair, add public key to GitHub, use git@ URLs",
    ])
    H(doc, "D. README that earns respect", 2)
    code(doc, [
        "Write a README with:",
        "## What this is",
        "## Features",
        "## Screenshots",
        "## Tech stack",
        "## Setup",
        "## Environment variables (names only)",
        "## Run",
        "## Demo URL",
        "## Author",
    ])
    H(doc, "E. Pull Requests (even solo)", 2)
    numbered(doc, [
        "Create branch: git checkout -b feature/readme-polish",
        "Commit changes",
        "Push branch: git push -u origin feature/readme-polish",
        "Open PR on GitHub comparing into main",
        "Write summary of what/why",
        "Optional: /code-review locally before merge",
        "Merge and pull main locally",
    ])
    H(doc, "F. Public repo safety audit", 2)
    bullets(doc, [
        "Search history for API keys if you ever committed secrets (rotation required)",
        "Confirm .env ignored",
        "Confirm node_modules/venv not tracked",
        "No private student data in sample DB dumps",
    ])
    pb(doc)

    # ========== WS10 full stack ==========
    H(doc, "Workshop 10 — Full-Stack Vertical Slice")
    p(doc, "A vertical slice is one thin feature implemented across frontend + backend + database end-to-end. Example: “Create task and see it in the list after refresh.” This beats building all frontend mockups first then all backend later.")
    H(doc, "Slice definition template", 2)
    code(doc, [
        "User story: As a student, I can add a task title and see it in my list after refresh.",
        "UI: input + add button + list",
        "API: POST /api/tasks, GET /api/tasks",
        "DB: insert row, select rows",
        "Acceptance:",
        "1) add 'Buy milk' appears in list",
        "2) refresh still shows it",
        "3) empty title rejected with message",
    ])
    H(doc, "Claude session pattern", 2)
    code(doc, [
        "/plan Implement ONLY the vertical slice above.",
        "List files you will touch. List test steps.",
        "No auth, no dark mode, no extras.",
    ])
    p(doc, "After approval, implement, run, test acceptance checklist, commit: “feat: add task create/list persistence slice”.")
    pb(doc)

    # ========== WS11 auth ==========
    H(doc, "Workshop 11 — Auth & Ownership")
    H(doc, "Concepts", 2)
    bullets(doc, [
        "Authentication: who are you?",
        "Authorization: what are you allowed to do?",
        "Password hashing: store irreversible transforms, never plain text",
        "Session/cookie or token: how the server remembers you after login",
        "Ownership checks: WHERE user_id = current_user on every private query",
    ])
    H(doc, "Test protocol (mandatory)", 2)
    numbered(doc, [
        "Register User A and User B",
        "Create data as A",
        "Login as B — must not see A’s private data",
        "Try to delete A’s item as B by ID if API allows guessing — must fail",
        "Logout clears access to protected pages",
    ])
    code(doc, [
        "/plan Add email/password auth with hashed passwords.",
        "Protect task routes so users only access own tasks.",
        "List security risks for beginners.",
        "Provide manual two-account test script.",
        "No OAuth unless already standard in our stack.",
    ])
    pb(doc)

    # ========== WS12 deploy ==========
    H(doc, "Workshop 12 — Deploy & Public Demo")
    H(doc, "Pre-deploy checklist", 2)
    bullets(doc, [
        "App runs clean locally",
        "README setup accurate",
        "Secrets only in env vars",
        "/security-review on the diff if available",
        "Choose host matching stack (static vs server vs container)",
    ])
    H(doc, "Production env vars", 2)
    p(doc, "Your laptop .env is not production. In the host dashboard, set the same variable NAMES with production VALUES (database URL, secret keys, public app URL). Never paste production secrets into GitHub issues.")
    H(doc, "Smoke test script", 2)
    numbered(doc, [
        "Open public URL in incognito",
        "Sign up fresh user",
        "Create two records",
        "Edit one, delete one",
        "Refresh; confirm persistence",
        "Phone test",
        "Put URL in README and commit",
    ])
    H(doc, "Demo script (5 minutes)", 2)
    table_data(doc, ["Time", "Content"], [
        ["30s", "Problem & user"],
        ["2–3 min", "Live happy path"],
        ["1 min", "3 decisions YOU made (scope, design, bugfix)"],
        ["30s", "What AI did well / next feature"],
    ])
    pb(doc)

    # ========== Capstone workshop ==========
    H(doc, "Workshop 13 — Capstone Integration Week")
    p(doc, "Combine everything: PRD → milestones → Claude plan/implement loops → GitHub → deploy → demo. Use FocusBoard (tasks) if your own idea is too large.")
    H(doc, "Capstone PRD fields (complete fully)", 2)
    bullets(doc, [
        "Name, problem, user, job-to-be-done",
        "MVP features (max 5)",
        "Out of scope",
        "Screens",
        "Data model",
        "Success demo script",
        "Risks (auth, data loss, time)",
    ])
    H(doc, "Suggested milestone backlog", 2)
    numbered(doc, [
        "M1 scaffold runs",
        "M2 data model + list/create",
        "M3 update/delete",
        "M4 polish UX",
        "M5 auth/ownership (if in scope)",
        "M6 README + GitHub",
        "M7 deploy + smoke test",
        "M8 demo rehearsal",
    ])
    pb(doc)

    # ========== Answer keys light ==========
    H(doc, "Appendix — Facilitator Notes for Workshops")
    bullets(doc, [
        "WS1 failure mode: PATH. Fix before teaching features.",
        "WS2: force written observations; do not let students click commands mindlessly.",
        "WS3–5: require tag/state maps so AI code becomes curriculum.",
        "WS6–7: one backend stack only for the room.",
        "WS8–9: secrets hygiene is a graded behavior, not a tip.",
        "WS11: two-account test is pass/fail gate.",
        "WS12: demo with offline video backup.",
    ])
    p(doc, "— End of Volume 2 — Deep Workshops —", bold=True)
    p(doc, "Use with: Vibe_Coding_Master_Teaching_Guide.docx + Vibe_Coding_Student_Study_Book.docx", size=9, italic=True, color=MUTED)

    doc.save(OUT)
    print("Wrote", OUT, "KB", round(os.path.getsize(OUT)/1024,1))

if __name__ == "__main__":
    build()
