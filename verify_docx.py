"""Verification helpers for the generated docx/pptx deliverables."""
import sys
from docx import Document


def heading_list(path, levels=(1, 2)):
    doc = Document(path)
    names = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3"}
    wanted = {names[level] for level in levels}
    return [p.text.strip() for p in doc.paragraphs if p.style.name in wanted and p.text.strip()]


def all_text(path):
    doc = Document(path)
    return "\n".join(p.text for p in doc.paragraphs)


def assert_contains(path, needles):
    text = all_text(path)
    missing = [n for n in needles if n not in text]
    assert not missing, f"{path} is missing expected text: {missing}"


if __name__ == "__main__":
    for h in heading_list(sys.argv[1]):
        print(h)
